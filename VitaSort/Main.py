import io
import json
import re
import shutil
import subprocess
from datetime import datetime
from typing import List, Dict, Optional, Tuple

import numpy as np
import pandas as pd
import streamlit as st

# Classic visuals deps
try:
    from PyPDF2 import PdfReader as _PdfReader
except ImportError:
    try:
        from pypdf import PdfReader as _PdfReader
    except ImportError:
        _PdfReader = None
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import matplotlib.pyplot as plt
from wordcloud import WordCloud

# Modern additions
try:
    import spacy
    SPACY_AVAILABLE = True
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        SPACY_AVAILABLE = False
        nlp = None
except ImportError:
    SPACY_AVAILABLE = False
    nlp = None

try:
    import textstat
    TEXTSTAT_AVAILABLE = True
except ImportError:
    TEXTSTAT_AVAILABLE = False

try:
    import faiss
    FAISS_AVAILABLE = True
except ImportError:
    FAISS_AVAILABLE = False

try:
    from langdetect import detect
    LANGDETECT_AVAILABLE = True
except ImportError:
    LANGDETECT_AVAILABLE = False


# App Config
st.set_page_config(page_title="VitaSort", page_icon="", layout="wide")
st.title("VitaSort")


# ========================= Legendary Utilities =========================
def pdf_to_text(data: bytes) -> Tuple[str, Optional[dict]]:
    meta = None
    text = ""
    
    # Try pdfminer first (best for text extraction)
    try:
        from pdfminer.high_level import extract_text
        with io.BytesIO(data) as fh:
            extracted = extract_text(fh)
            if extracted and len(extracted.strip()) > 10:
                text = extracted
    except Exception:
        pass
    
    # Try pypdf/PyPDF2 as fallback or supplement
    try:
        from pypdf import PdfReader
        with io.BytesIO(data) as fh:
            reader = PdfReader(fh)
            
            # If no text from pdfminer, extract from pypdf
            if not text or len(text.strip()) < 50:
                parts = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        parts.append(page_text)
                if parts:
                    text = "\n\n".join(parts)
            
            # Extract metadata
            try:
                md = reader.metadata
                if md:
                    meta = {k: str(v) for k, v in md.items()}
            except Exception:
                pass
                
    except Exception:
        pass
    
    # Clean up extracted text
    if text:
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        # Remove control characters except newlines and tabs
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
    
    return text.strip(), meta


def docx_to_text(data: bytes) -> str:
    try:
        import docx
        with io.BytesIO(data) as fh:
            d = docx.Document(fh)
            text_parts = []
            
            # Extract paragraphs
            for p in d.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text)
            
            # Extract text from tables
            for table in d.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            result = "\n".join(text_parts)
            if result:
                return result
    except Exception:
        pass
    
    # Fallback to docx2txt
    try:
        import docx2txt  # type: ignore
        with io.BytesIO(data) as fh:
            result = docx2txt.process(fh)
            if result:
                return result
    except Exception:
        pass
    
    return ""


def rtf_to_text(data: bytes) -> str:
    try:
        from striprtf.striprtf import rtf_to_text

        return rtf_to_text(data.decode(errors="ignore"))
    except Exception:
        return ""


EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
PHONE_RE = re.compile(r"(?:\+?\d{1,4}[\s.-]?)?(?:\(?\d{1,4}\)?[\s.-]?)?\d{1,4}[\s.-]?\d{1,4}[\s.-]?\d{1,9}")
LINKEDIN_RE = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9_-]+/?|linkedin\.com/in/[A-Za-z0-9_-]+", re.I)
GITHUB_RE = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9_-]+/?|github\.com/[A-Za-z0-9_-]+", re.I)
PORTFOLIO_RE = re.compile(r"(?:portfolio|website|blog)\s*:?\s*(https?://[^\s]+)", re.I)
URL_RE = re.compile(r"https?://[^\s<>\"]+", re.I)
YEAR_RANGE_RE = re.compile(r"(\b\d{4}\b)\s*[-–—to]\s*(\b\d{4}\b|Present|Current|Now|Till date|Till now|Ongoing)", re.I)
MONTH_YEAR_RE = re.compile(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s*(\d{4})", re.I)
YEAR_RE = re.compile(r"\b(19\d{2}|20\d{2})\b")


SKILLS = {
    "programming": ["python", "java", "javascript", "typescript", "c++", "c#", "go", "rust", "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "perl", "shell", "bash", "powershell"],
    "data": ["sql", "nosql", "postgres", "postgresql", "mysql", "mongodb", "redis", "cassandra", "elasticsearch", "data engineering", "etl", "spark", "hadoop", "kafka", "airflow", "databricks", "snowflake", "bigquery"],
    "ml_ai": ["machine learning", "deep learning", "nlp", "natural language processing", "computer vision", "scikit-learn", "pytorch", "tensorflow", "keras", "transformers", "hugging face", "opencv", "yolo", "bert", "gpt", "llm", "generative ai"],
    "analytics": ["excel", "powerbi", "power bi", "tableau", "looker", "analytics", "statistics", "data analysis", "data visualization", "pandas", "numpy", "matplotlib", "seaborn", "plotly"],
    "cloud": ["aws", "amazon web services", "azure", "microsoft azure", "gcp", "google cloud", "kubernetes", "k8s", "docker", "terraform", "ansible", "jenkins", "ci/cd", "serverless", "lambda", "ec2", "s3"],
    "web": ["react", "reactjs", "vue", "vuejs", "angular", "django", "flask", "fastapi", "node", "nodejs", "express", "next.js", "nuxt", "svelte", "html", "css", "sass", "tailwind", "bootstrap"],
    "mobile": ["android", "ios", "react native", "flutter", "swift", "kotlin", "xamarin", "cordova", "ionic"],
    "devops": ["devops", "ci/cd", "jenkins", "gitlab", "github actions", "circleci", "travis", "docker", "kubernetes", "terraform", "ansible", "chef", "puppet", "monitoring", "prometheus", "grafana"],
    "testing": ["testing", "unit testing", "integration testing", "jest", "pytest", "selenium", "cypress", "junit", "testng", "qa", "quality assurance", "automation"],
    "database": ["database", "dbms", "sql", "nosql", "oracle", "mssql", "sqlite", "dynamodb", "couchdb", "neo4j", "graph database"],
    "security": ["security", "cybersecurity", "penetration testing", "ethical hacking", "owasp", "encryption", "ssl", "tls", "authentication", "oauth", "jwt", "firewall"],
    "soft": ["communication", "leadership", "teamwork", "team work", "problem solving", "critical thinking", "presentation", "agile", "scrum", "project management", "time management", "collaboration", "mentoring", "coaching"],
}


def extract_contacts(text: str) -> Dict[str, Optional[str]]:
    """ENHANCED contact extraction with better name detection and multiple fallback strategies"""
    emails = EMAIL_RE.findall(text)
    email = emails[0] if emails else None
    
    phones = PHONE_RE.findall(text)
    # Filter out very short phone numbers (likely false positives)
    valid_phones = [p for p in phones if len(re.sub(r'[^\d]', '', p)) >= 7]
    phone = valid_phones[0] if valid_phones else None
    
    name = None
    lines = text.splitlines()
    
    # Try multiple name extraction strategies - ENHANCED
    for i, line in enumerate(lines[:20]):  # Check first 20 lines (increased from 15)
        s = line.strip()
        if not s or len(s) < 3 or len(s) > 100:
            continue
            
        words = s.split()
        if not (2 <= len(words) <= 6):  # Allow up to 6 words for full names with titles
            continue
        
        # Strategy 1: Look for capitalized words (typical name format)
        caps = sum(1 for w in words if w and w[0].isupper())
        if caps >= max(2, len(words) - 1):
            # Avoid email, phone, URLs, common headers in name
            if not any(x in s.lower() for x in ['@', 'http', '.com', '.net', '.org', 'phone', 'email', 'resume', 'cv', 'curriculum', 'vitae', 'address', 'linkedin', 'github', 'portfolio']):
                # Avoid lines with too many numbers (addresses, phone numbers)
                if sum(c.isdigit() for c in s) < len(s) * 0.3:
                    name = s
                    break
        
        # Strategy 2: Look for explicit name labels
        if any(keyword in s.lower() for keyword in ['name:', 'full name:', 'candidate:', 'applicant:', 'personal details']):
            name_part = re.sub(r'(?i)(full\s+)?name\s*:?\s*|candidate\s*:?\s*|applicant\s*:?\s*|personal\s+details\s*:?\s*', '', s).strip()
            if 3 <= len(name_part) <= 80:
                name = name_part
                break
        
        # Strategy 3: First line if it looks like a name (2-4 capitalized words, no special chars)
        if i == 0 and 2 <= len(words) <= 4:
            if all(w[0].isupper() for w in words if w) and not re.search(r'[^\w\s\.]', s):
                name = s
                break
    
    # If still no name found, try to extract from email
    if not name and email:
        username = email.split('@')[0]
        # Convert common email formats to names
        name_guess = username.replace('.', ' ').replace('_', ' ').replace('-', ' ')
        # Only use if it looks like a real name (2+ parts, reasonable length)
        parts = name_guess.split()
        if len(parts) >= 2 and 3 <= len(name_guess) <= 50:
            name = name_guess.title()
    
    return {"name": name, "email": email, "phone": phone}


def extract_social(text: str) -> Dict[str, List[str]]:
    """Extract ALL URLs and categorize by platform - works with any custom platform"""
    # Extract all URLs (with or without http/https)
    all_urls = list(set(URL_RE.findall(text)))
    
    # Also find URLs without protocol
    url_pattern_no_protocol = re.compile(r'\b(?:[a-z0-9-]+\.)+[a-z]{2,}(?:/[^\s]*)?', re.I)
    urls_no_protocol = [url for url in url_pattern_no_protocol.findall(text) 
                        if '.' in url and not url.endswith('.') 
                        and len(url) > 5 and not url[0].isdigit()]
    
    # Combine and deduplicate
    all_urls_combined = list(set(all_urls + ['https://' + u if not u.startswith('http') else u for u in urls_no_protocol]))
    
    # Platform categorization with flexible matching
    linkedin = [url for url in all_urls_combined if 'linkedin.com' in url.lower()]
    github = [url for url in all_urls_combined if 'github.com' in url.lower() or 'gitlab.com' in url.lower()]
    twitter = [url for url in all_urls_combined if 'twitter.com' in url.lower() or 'x.com' in url.lower()]
    stackoverflow = [url for url in all_urls_combined if 'stackoverflow.com' in url.lower()]
    medium = [url for url in all_urls_combined if 'medium.com' in url.lower()]
    behance = [url for url in all_urls_combined if 'behance.net' in url.lower()]
    dribbble = [url for url in all_urls_combined if 'dribbble.com' in url.lower()]
    
    # Detect portfolio/personal sites (common indicators)
    portfolio_keywords = ['portfolio', 'website', 'blog', 'personal']
    portfolio = [url for url in all_urls_combined if any(kw in url.lower() for kw in portfolio_keywords)]
    
    # Custom platforms - anything not in known categories
    known_platforms = set(linkedin + github + twitter + stackoverflow + medium + behance + dribbble + portfolio)
    custom = [url for url in all_urls_combined if url not in known_platforms]
    
    return {
        "linkedin": linkedin,
        "github": github,
        "portfolio": portfolio,
        "twitter": twitter,
        "stackoverflow": stackoverflow,
        "medium": medium,
        "behance": behance,
        "dribbble": dribbble,
        "custom": custom,  # NEW: Any other platform URLs
        "all": all_urls_combined
    }


def extract_education(text: str) -> List[dict]:
    out = []
    # Expanded degree patterns
    degree_re = re.compile(
        r"(Bachelor|Master|BSc|MSc|BA|MA|PhD|MBA|BS|MS|MPhil|B\.Tech|M\.Tech|BEng|MEng|B\.E\.|M\.E\.|" +
        r"Associate|Diploma|Certificate|Doctorate|DBA|JD|MD|BBA|MCA|BCA|B\.Sc|M\.Sc|B\.A\.|M\.A\.|" +
        r"High School|Secondary|Higher Secondary|HSC|SSC|GCSE|A-Level|O-Level|Degree)", 
        re.I
    )
    
    # Education section markers
    edu_section_re = re.compile(r"^(education|academic|qualification|degree|university|college)\s*:?", re.I)
    
    lines = text.splitlines()
    in_edu_section = False
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        # Check if we're entering education section
        if edu_section_re.match(stripped):
            in_edu_section = True
            continue
        
        # Check for section end
        if in_edu_section and stripped and any(marker in stripped.lower() for marker in ['experience', 'work history', 'employment', 'skills', 'projects']):
            in_edu_section = False
        
        # Look for degree mentions
        if degree_re.search(stripped):
            entry = {"raw": stripped}
            
            # Extract degree name
            degree_match = degree_re.search(stripped)
            if degree_match:
                entry["degree"] = degree_match.group(1)
            
            # Extract years/graduation date
            years = YEAR_RE.findall(stripped)
            if years:
                entry["years"] = years
                entry["graduation_year"] = years[-1]  # Last year is usually graduation
            
            # Look for month-year formats
            month_years = MONTH_YEAR_RE.findall(stripped)
            if month_years:
                entry["dates"] = month_years
            
            # Extract institution from same line or next lines
            institution = None
            
            # Try to extract from same line (after degree)
            remaining = degree_re.sub('', stripped).strip()
            if len(remaining) > 5:
                # Remove years and common words
                for year in years:
                    remaining = remaining.replace(year, '')
                remaining = re.sub(r'\b(in|from|at|\d{4}|[,.-])\b', ' ', remaining, flags=re.I).strip()
                if len(remaining) > 3:
                    institution = remaining
            
            # Try next 2 lines for institution
            if not institution:
                for j in range(i + 1, min(i + 3, len(lines))):
                    nxt = lines[j].strip()
                    # Skip if it's another degree or section header
                    if degree_re.search(nxt) or edu_section_re.match(nxt):
                        break
                    if 5 < len(nxt) < 150 and not nxt.startswith(('-', '•', '*')):
                        institution = nxt
                        break
            
            if institution:
                entry["institution"] = institution
            
            out.append(entry)
    
    return out


def extract_experience(text: str) -> Tuple[List[dict], float, List[Tuple[datetime, datetime]]]:
    timeline = []
    date_ranges: List[Tuple[datetime, datetime]] = []
    
    # Common job title keywords
    job_title_keywords = ['engineer', 'developer', 'manager', 'analyst', 'specialist', 'consultant', 
                          'architect', 'lead', 'senior', 'junior', 'intern', 'associate', 'director',
                          'coordinator', 'supervisor', 'executive', 'officer', 'administrator', 'designer']
    
    lines = text.splitlines()
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        # Look for year ranges
        m = YEAR_RANGE_RE.search(stripped)
        if not m:
            continue
        
        try:
            start = int(m.group(1))
            end_raw = m.group(2).strip()
            
            # Parse end date
            if end_raw.lower() in ("present", "current", "now", "ongoing", "till date", "till now"):
                end = datetime.now().year
            else:
                end = int(end_raw)
            
            # Validate years
            if not (1970 <= start <= datetime.now().year and 1970 <= end <= datetime.now().year + 1):
                continue
            
            if start > end:
                continue
            
            sdt = datetime(start, 1, 1)
            edt = datetime(end, 12, 31)
            date_ranges.append((sdt, edt))
            
            # Try to extract job title and company
            job_title = None
            company = None
            
            # Look in same line
            line_without_dates = YEAR_RANGE_RE.sub('', stripped).strip()
            line_without_dates = re.sub(r'[|•●◦▪▫-]', '', line_without_dates).strip()
            
            if any(keyword in line_without_dates.lower() for keyword in job_title_keywords):
                job_title = line_without_dates
            
            # Look for company in next line
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                # If next line doesn't have dates and is reasonable length, it might be company
                if not YEAR_RANGE_RE.search(next_line) and 3 < len(next_line) < 100:
                    company = next_line
            
            # Look for job title in previous line
            if not job_title and i > 0:
                prev_line = lines[i - 1].strip()
                if any(keyword in prev_line.lower() for keyword in job_title_keywords) and len(prev_line) < 100:
                    job_title = prev_line
            
            entry = {
                "period": m.group(0), 
                "raw": stripped,
                "start_year": start,
                "end_year": end,
                "duration_years": round((edt - sdt).days / 365.25, 1)
            }
            
            if job_title:
                entry["title"] = job_title
            if company:
                entry["company"] = company
            
            timeline.append(entry)
            
        except (ValueError, AttributeError):
            # Skip malformed dates
            continue
    
    # Calculate total years of experience
    years = sum(max(0, (e - s).days) / 365.25 for s, e in date_ranges)
    
    # Detect employment gaps (6+ months)
    gaps = []
    if len(date_ranges) >= 2:
        rs = sorted(date_ranges, key=lambda x: x[0])
        for (s1, e1), (s2, e2) in zip(rs, rs[1:]):
            gap_days = (s2 - e1).days
            if gap_days > 183:  # More than 6 months
                gaps.append((e1, s2))
    
    return timeline, years, gaps


def tag_skills(text: str) -> Dict[str, List[str]]:
    """ULTRA-ENHANCED skill tagging with 10x more dynamic detection - captures technologies, frameworks, methodologies, and context"""
    found = {cat: [] for cat in SKILLS}
    found["dynamic_skills"] = []
    found["frameworks"] = []
    found["methodologies"] = []
    found["tools"] = []
    found["platforms"] = []
    found["certifications"] = []
    
    text_lower = text.lower()
    lines = text.split('\n')
    
    # 1. Fixed dictionary matching (original)
    for category, skills_list in SKILLS.items():
        for skill in skills_list:
            pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(pattern, text_lower):
                found[category].append(skill)
    
    # 2. MASSIVELY ENHANCED DYNAMIC EXTRACTION
    
    # Acronyms (2-8 uppercase letters) - expanded range
    acronyms = re.findall(r'\b[A-Z]{2,8}\b', text)
    exclude_acronyms = ['CV', 'USA', 'UK', 'MBA', 'BSC', 'MSC', 'PHD', 'CEO', 'CTO', 'CFO', 'VP', 'HR', 'IT', 'PM', 'QA', 'BA', 'MA']
    tech_acronyms = [a for a in acronyms if a not in exclude_acronyms]
    
    # CamelCase technologies (e.g., TensorFlow, PyTorch, GitHub, MongoDB)
    camel_case = re.findall(r'\b[A-Z][a-z]+[A-Z][a-zA-Z]*\b', text)
    
    # Versioned technologies (e.g., Python 3.9, Node.js 16, Java 11, .NET 6)
    versioned = re.findall(r'\b([A-Za-z\.]+[\w]*)\s*[v]?[\d]+(?:\.[\d]+)*\b', text)
    
    # Hyphenated tools (e.g., scikit-learn, test-driven, cross-platform)
    hyphenated = re.findall(r'\b[a-z]+(?:-[a-z]+)+\b', text_lower)
    
    # Dotted technologies (e.g., Node.js, Vue.js, ASP.NET, D3.js)
    dotted = re.findall(r'\b[A-Za-z]+\.[A-Za-z]+\b', text)
    
    # Programming patterns (e.g., "in Python", "using Java", "with React")
    context_skills = []
    for pattern in [r'\b(?:in|using|with|via|on)\s+([A-Z][a-z\w]+)\b', r'\b([A-Z][a-z\w]+)\s+(?:development|programming|coding|scripting)\b']:
        context_skills.extend(re.findall(pattern, text))
    
    # Framework indicators (e.g., "Django framework", "React library", "Express.js")
    frameworks = re.findall(r'\b([A-Z][a-z\w\.]+)\s+(?:framework|library|toolkit|platform|engine)\b', text, re.IGNORECASE)
    found["frameworks"] = list(set(frameworks))[:15]
    
    # Methodology indicators (Agile, Scrum, TDD, CI/CD, DevOps practices)
    methodologies = re.findall(r'\b(Agile|Scrum|Kanban|Waterfall|TDD|BDD|CI/CD|DevOps|GitFlow|Microservices|Serverless|RESTful|GraphQL|SOLID|DRY|KISS|LEAN|Six Sigma|ITIL|SAFe)\b', text, re.IGNORECASE)
    found["methodologies"] = list(set([m.upper() if len(m) <= 4 else m.title() for m in methodologies]))
    
    # Tool/Software patterns ("Proficient in X", "Expert with Y", "Experience using Z")
    proficiency_tools = []
    for pattern in [r'(?:proficient|expert|experienced|skilled)\s+(?:in|with|using)\s+([A-Z][a-z\w\.\-]+)', r'(?:knowledge|experience)\s+(?:of|in|with)\s+([A-Z][a-z\w\.\-]+)']:
        proficiency_tools.extend(re.findall(pattern, text, re.IGNORECASE))
    found["tools"] = list(set(proficiency_tools))[:15]
    
    # Cloud/Database services (AWS S3, Azure Blob, GCP BigQuery, AWS Lambda)
    cloud_services = re.findall(r'\b(AWS|Azure|GCP|Google Cloud)\s+([A-Z][a-z\w]+)\b', text)
    cloud_tech = [f"{svc[0]} {svc[1]}" for svc in cloud_services]
    
    # Operating Systems & Platforms
    os_platforms = re.findall(r'\b(Linux|Unix|Windows|macOS|iOS|Android|Ubuntu|CentOS|RedHat|Debian|Fedora|Arch|FreeBSD|Solaris)\b', text, re.IGNORECASE)
    found["platforms"] = list(set([s.title() for s in os_platforms]))
    
    # Build & Deployment Tools
    build_tools = re.findall(r'\b(Jenkins|CircleCI|TravisCI|GitLab CI|GitHub Actions|Bamboo|TeamCity|Gradle|Maven|Webpack|Gulp|npm|yarn|pip|Terraform|Ansible|Chef|Puppet|CloudFormation)\b', text, re.IGNORECASE)
    
    # Databases (including NoSQL and NewSQL)
    databases = re.findall(r'\b(MongoDB|PostgreSQL|MySQL|Oracle|SQL Server|Redis|Cassandra|DynamoDB|Elasticsearch|Neo4j|CouchDB|MariaDB|SQLite|Firestore|Snowflake|BigQuery|Redshift)\b', text, re.IGNORECASE)
    
    # Design & Testing Tools
    design_test = re.findall(r'\b(Figma|Sketch|Adobe XD|InVision|Photoshop|Illustrator|Selenium|Cypress|Jest|Mocha|Pytest|JUnit|TestNG|Postman|SoapUI|LoadRunner|JMeter)\b', text, re.IGNORECASE)
    
    # Container & Orchestration
    containers = re.findall(r'\b(Docker|Kubernetes|K8s|Helm|Rancher|OpenShift|Podman|Nomad|Compose|Swarm)\b', text, re.IGNORECASE)
    
    # Monitoring & Logging
    monitoring = re.findall(r'\b(Grafana|Prometheus|ELK|Kibana|Logstash|Datadog|New Relic|Splunk|CloudWatch|AppDynamics|Nagios|Zabbix)\b', text, re.IGNORECASE)
    
    # Message Queues & Streaming
    messaging = re.findall(r'\b(Kafka|RabbitMQ|Redis|ActiveMQ|ZeroMQ|Pulsar|MQTT|NATS)\b', text, re.IGNORECASE)
    
    # API & Web Services
    api_tools = re.findall(r'\b(REST|GraphQL|gRPC|SOAP|Swagger|OpenAPI|Postman|Insomnia)\b', text, re.IGNORECASE)
    
    # Version Control & Collaboration
    vcs_tools = re.findall(r'\b(Git|GitHub|GitLab|Bitbucket|SVN|Mercurial|Perforce|Jira|Confluence|Slack|Teams)\b', text, re.IGNORECASE)
    
    # Skills from bullet points (lines starting with bullet characters)
    bullet_skills = []
    for line in lines:
        if re.match(r'^\s*[•●○▪▫■□‣⁃\\-\\*]', line):
            # Extract capitalized words/phrases from bullet points
            caps = re.findall(r'\b[A-Z][a-z\w\\.\\-]*(?:\s+[A-Z][a-z\w\\.\\-]*)*\b', line)
            bullet_skills.extend([c for c in caps if len(c) > 2 and c not in exclude_acronyms])
    
    # Skills from "Skills:" section lines
    skill_section_items = []
    in_skills_section = False
    for line in lines:
        if re.search(r'^\s*(?:Technical\s+)?Skills?\s*:?\s*$', line, re.IGNORECASE):
            in_skills_section = True
            continue
        if in_skills_section:
            if re.match(r'^\s*(?:Education|Experience|Projects|Certifications|Work History|Employment)', line, re.IGNORECASE):
                break
            # Extract from comma/pipe/semicolon separated lists
            items = re.split(r'[,|;]', line)
            for item in items:
                cleaned = item.strip(' •●○▪▫■□‣⁃\\-\\*\\t')
                if cleaned and len(cleaned) > 2 and not re.match(r'^\d+$', cleaned):
                    skill_section_items.append(cleaned)
    
    # Programming languages with context (e.g., "3 years of Python", "proficient in Java")
    lang_context = re.findall(r'(?:\d+\s+years?\s+(?:of|in|with)\s+|proficient\s+in\s+|expert\s+in\s+|experience\s+with\s+)([A-Z][a-z\w\+\#]+)', text, re.IGNORECASE)
    
    # Certifications embedded in skills
    cert_patterns = re.findall(r'\b(AWS Certified|Azure Certified|Google Cloud Certified|Oracle Certified|Microsoft Certified|Cisco Certified|CompTIA|PMP|CISSP|CEH|CKA|CKAD)\b', text, re.IGNORECASE)
    found["certifications"] = list(set([c.title() for c in cert_patterns]))[:10]
    
    # Combine all dynamic skills
    all_dynamic = (
        tech_acronyms + camel_case + [v for v in versioned if v] + 
        hyphenated + dotted + context_skills + cloud_tech + 
        [s.title() for s in build_tools] + 
        [s.title() for s in databases] + 
        [s.title() for s in design_test] + 
        [s.title() for s in containers] + 
        [s.title() for s in monitoring] + 
        [s.title() for s in messaging] + 
        [s.upper() if len(s) <= 4 else s.title() for s in api_tools] +
        [s.title() for s in vcs_tools] +
        bullet_skills[:30] +  # Limit bullet skills
        skill_section_items[:40] +  # Limit section skills
        [s.title() for s in lang_context]
    )
    
    # Deduplicate and filter
    seen = set()
    unique_dynamic = []
    common_words = {'and', 'the', 'for', 'with', 'from', 'have', 'this', 'that', 'will', 'your', 'can', 'all', 'not', 'but', 'are', 'was', 'were', 'been', 'has', 'had', 'which', 'also', 'their', 'there', 'where', 'when', 'what', 'who', 'how', 'why', 'more', 'most', 'some', 'such', 'into', 'through', 'during', 'before', 'after', 'above', 'below', 'between', 'under', 'over'}
    
    for skill in all_dynamic:
        skill_clean = str(skill).strip()
        skill_lower = skill_clean.lower()
        # Filter out common words and very short items
        if (len(skill_clean) >= 2 and 
            skill_lower not in common_words and
            skill_clean not in seen and
            not skill_clean.isdigit()):
            seen.add(skill_clean)
            unique_dynamic.append(skill_clean)
    
    found["dynamic_skills"] = unique_dynamic[:60]  # Increased limit to 60
    
    return found


def cosine(a, b) -> float:
    from numpy import dot
    from numpy.linalg import norm

    d = norm(a) * norm(b)
    return float(dot(a, b) / d) if d else 0.0


def embed_texts(texts: List[str]) -> np.ndarray:
    try:
        from sentence_transformers import SentenceTransformer

        model = SentenceTransformer("all-MiniLM-L6-v2")
        return np.array(model.encode(texts))
    except Exception:
        X = TfidfVectorizer().fit_transform(texts)
        return X.toarray()


def semantic_and_hard(jd: str, cv: str) -> Tuple[float, float]:
    mat = embed_texts([jd, cv])
    sem = cosine(mat[0], mat[1])
    jd_words = set(re.findall(r"\b\w+\b", jd.lower()))
    cv_words = set(re.findall(r"\b\w+\b", cv.lower()))
    inter = len([w for w in jd_words if w in cv_words and len(w) > 2])
    hard = min(1.0, inter / max(10, len(jd_words)))
    return float(sem), float(hard)


def scoring_breakdown(jd: str, cv: str, semantic_score: float, hard_skill_score: float, total_years: float,
                      soft_tags_count: int, social: Dict[str, List[str]], contacts: Dict[str, Optional[str]],
                      education: List[dict], experience_points: int, custom_weights: Optional[Dict] = None) -> Dict:
    jd_fit = (semantic_score * 0.7) + (hard_skill_score * 0.3)
    hard_skills = hard_skill_score
    experience_quality = min(1.0, total_years / 10.0)
    soft_skills = min(1.0, soft_tags_count / 8.0)
    digital_presence = min(1.0, (len(social.get("linkedin", [])) + len(social.get("github", []))) / 2.0)
    formatting = 0.0
    if contacts.get("email"):
        formatting += 0.2
    if education:
        formatting += 0.2
    if experience_points:
        formatting += 0.2
    if len(cv) > 400:
        formatting += 0.2
    if "skills" in cv.lower():
        formatting += 0.2
    formatting = min(1.0, formatting)
    
    # Apply custom weights if provided
    if custom_weights:
        overall = (custom_weights["hard"] * hard_skills + 
                  custom_weights["exp"] * experience_quality + 
                  custom_weights["soft"] * soft_skills + 
                  custom_weights["digital"] * digital_presence + 
                  custom_weights["format"] * formatting)
    else:
        overall = 0.40 * hard_skills + 0.30 * experience_quality + 0.10 * soft_skills + 0.10 * digital_presence + 0.10 * formatting
    
    return {"jd_fit": round(jd_fit, 4), "overall": round(overall, 4),
            "breakdown": {"hard_skills": round(hard_skills, 4), "experience": round(experience_quality, 4),
                           "soft_skills": round(soft_skills, 4), "digital_presence": round(digital_presence, 4),
                           "formatting": round(formatting, 4), "semantic": round(semantic_score, 4)}}


def email_osint(email: Optional[str]) -> Optional[dict]:
    if not email:
        return None
    out = {"valid_syntax": False, "holehe": None}
    try:
        from email_validator import validate_email, EmailNotValidError

        try:
            validate_email(email)
            out["valid_syntax"] = True
        except EmailNotValidError:
            out["valid_syntax"] = False
    except Exception:
        pass
    if shutil.which("holehe") and out["valid_syntax"]:
        try:
            res = subprocess.run(["holehe", "-j", email], capture_output=True, text=True, timeout=30)
            if res.returncode == 0 and res.stdout.strip():
                out["holehe"] = json.loads(res.stdout.strip())
        except Exception:
            pass
    return out


def fact_check_name(name: Optional[str]) -> List[str]:
    if not name:
        return []
    try:
        from googlesearch import search  # type: ignore

        return [url for url in search(name, num_results=5)]
    except Exception:
        return []


def full_text_search(text: str, query: str) -> bool:
    return query.lower() in text.lower()


def boolean_search(text: str, query: str) -> bool:
    tokens = re.findall(r'\w+|\(|\)|AND|OR|NOT', query, flags=re.I)
    expr = []
    s = text.lower()
    for tok in tokens:
        up = tok.upper()
        if up in ("AND", "OR", "NOT", "(", ")"):
            expr.append(up)
        else:
            expr.append(str(tok.lower() in s))
    try:
        py = " ".join(expr).replace("AND", "and").replace("OR", "or").replace("NOT", "not")
        return bool(eval(py))
    except Exception:
        return False


def deduplicate(profiles: List[dict]) -> List[dict]:
    seen = set()
    out = []
    for p in profiles:
        key = p.get("contacts", {}).get("email") or p.get("contacts", {}).get("phone")
        if not key or key not in seen:
            out.append(p)
            if key:
                seen.add(key)
    return out


# ========================= Advanced Features =========================
def extract_location(text: str) -> Optional[str]:
    """Extract location from CV using regex patterns"""
    patterns = [
        r"(?:Location|Address|City):\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*(?:,\s*[A-Z]{2})?)",
        r"([A-Z][a-z]+,\s*[A-Z]{2,3}\s+\d{5})",
        r"([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*(?:USA|UK|Canada|India))"
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1)
    return None


def calculate_location_distance(loc1: str, loc2: str) -> Optional[float]:
    """Calculate distance between two locations in km"""
    try:
        from geopy.geocoders import Nominatim
        from geopy.distance import geodesic
        geolocator = Nominatim(user_agent="vitasort")
        l1 = geolocator.geocode(loc1, timeout=5)
        l2 = geolocator.geocode(loc2, timeout=5)
        if l1 and l2:
            return geodesic((l1.latitude, l1.longitude), (l2.latitude, l2.longitude)).km
    except Exception:
        pass
    return None


def detect_skill_proficiency(text: str, skill: str) -> str:
    """Detect skill proficiency level based on context"""
    lower_text = text.lower()
    if skill.lower() not in lower_text:
        return "None"
    
    senior_indicators = ["senior", "expert", "advanced", "lead", "architect", "principal", "5+ years", "8+ years"]
    mid_indicators = ["mid-level", "intermediate", "proficient", "3 years", "4 years"]
    junior_indicators = ["junior", "beginner", "entry", "learning", "familiar with", "1 year", "2 years"]
    
    skill_context = " ".join([line for line in lower_text.split('\n') if skill.lower() in line])
    
    if any(ind in skill_context for ind in senior_indicators):
        return "Senior"
    elif any(ind in skill_context for ind in mid_indicators):
        return "Mid-Level"
    elif any(ind in skill_context for ind in junior_indicators):
        return "Junior"
    return "Proficient"


def generate_career_summary(exp_list: List[dict], years: float) -> str:
    """Auto-generate career highlight summary"""
    if not exp_list:
        return "No experience data available."
    
    summary_parts = []
    summary_parts.append(f"Total Experience: {years:.1f} years")
    
    roles = []
    for exp in exp_list:
        raw = exp.get("raw", "").lower()
        if "senior" in raw or "lead" in raw:
            roles.append("Senior-level")
        elif "manager" in raw or "director" in raw:
            roles.append("Management")
        elif "engineer" in raw or "developer" in raw:
            roles.append("Technical")
    
    if roles:
        summary_parts.append(f"Roles: {', '.join(set(roles))}")
    
    return " | ".join(summary_parts)


def estimate_salary(years: float, location: Optional[str], skills: Dict) -> Dict:
    """Estimate salary range based on experience, location, and skills"""
    base_salary = 50000
    per_year = 5000
    
    estimated = base_salary + (years * per_year)
    
    # Location multiplier
    if location:
        loc_lower = location.lower()
        if any(city in loc_lower for city in ["san francisco", "new york", "seattle", "boston"]):
            estimated *= 1.5
        elif any(city in loc_lower for city in ["austin", "denver", "chicago", "los angeles"]):
            estimated *= 1.3
    
    # Skills premium
    premium_skills = len(skills.get("ml_ai", [])) * 10000 + len(skills.get("cloud", [])) * 5000
    estimated += premium_skills
    
    return {
        "min": int(estimated * 0.85),
        "max": int(estimated * 1.15),
        "estimated": int(estimated)
    }


def check_github_activity(github_url: str) -> Optional[Dict]:
    """Scan GitHub profile for activity metrics"""
    try:
        import requests
        username = github_url.rstrip('/').split('/')[-1]
        resp = requests.get(f"https://api.github.com/users/{username}", timeout=5)
        if resp.status_code == 200:
            data = resp.json()
            return {
                "public_repos": data.get("public_repos", 0),
                "followers": data.get("followers", 0),
                "created_at": data.get("created_at", "")
            }
    except Exception:
        pass
    return None


def check_hibp(email: str) -> Optional[Dict]:
    """Check HaveIBeenPwned for email breaches"""
    try:
        import requests
        resp = requests.get(f"https://haveibeenpwned.com/api/v3/breachedaccount/{email}", 
                          headers={"User-Agent": "VitaSort"}, timeout=5)
        if resp.status_code == 200:
            return {"breached": True, "count": len(resp.json())}
        elif resp.status_code == 404:
            return {"breached": False, "count": 0}
    except Exception:
        pass
    return None


def detect_job_hopping(timeline: List[dict]) -> Dict:
    """Analyze job hopping patterns"""
    if len(timeline) < 2:
        return {"hopper": False, "avg_tenure": 0}
    
    tenures = []
    for exp in timeline:
        period_str = exp.get("period", "")
        match = YEAR_RANGE_RE.search(period_str)
        if match:
            start = int(match.group(1))
            end_str = match.group(2)
            end = datetime.now().year if end_str.lower() in ("present", "current") else int(end_str)
            tenures.append(end - start)
    
    if tenures:
        avg = sum(tenures) / len(tenures)
        return {"hopper": avg < 2, "avg_tenure": round(avg, 1), "job_count": len(timeline)}
    return {"hopper": False, "avg_tenure": 0, "job_count": len(timeline)}


def blind_screening_filter(profile: Dict) -> Dict:
    """Remove bias attributes for blind screening"""
    blind = profile.copy()
    blind["contacts"]["name"] = "Candidate"
    blind["contacts"]["email"] = "***@***.com"
    blind["contacts"]["phone"] = "***-***-****"
    blind["fact_links"] = []
    return blind


def detect_text_similarity(text1: str, text2: str) -> float:
    """Detect plagiarism/similarity between two texts"""
    try:
        vec = TfidfVectorizer().fit_transform([text1, text2])
        return float(cosine_similarity(vec[0:1], vec[1:2])[0][0])
    except Exception:
        return 0.0


# ========================= Dashboard Analytics =========================
def create_dashboard_metrics(results: List[dict]) -> Dict:
    """Generate comprehensive dashboard metrics"""
    if not results:
        return {}
    
    total_candidates = len(results)
    avg_score = np.mean([r["score"]["overall"] for r in results])
    avg_years = np.mean([r["years_experience"] for r in results])
    
    with_linkedin = sum(1 for r in results if r["social"].get("linkedin"))
    with_github = sum(1 for r in results if r["social"].get("github"))
    with_gaps = sum(1 for r in results if r.get("gaps"))
    
    top_skills = {}
    for r in results:
        for cat, skills in r["tags"].items():
            for skill in skills:
                top_skills[skill] = top_skills.get(skill, 0) + 1
    
    top_10_skills = sorted(top_skills.items(), key=lambda x: x[1], reverse=True)[:10]
    
    return {
        "total": total_candidates,
        "avg_score": round(avg_score, 2),
        "avg_years": round(avg_years, 1),
        "linkedin_pct": round((with_linkedin / total_candidates) * 100, 1),
        "github_pct": round((with_github / total_candidates) * 100, 1),
        "gaps_pct": round((with_gaps / total_candidates) * 100, 1),
        "top_skills": top_10_skills,
        "score_distribution": {
            "excellent": sum(1 for r in results if r["score"]["overall"] >= 0.8),
            "good": sum(1 for r in results if 0.6 <= r["score"]["overall"] < 0.8),
            "average": sum(1 for r in results if 0.4 <= r["score"]["overall"] < 0.6),
            "below": sum(1 for r in results if r["score"]["overall"] < 0.4)
        }
    }


def create_score_distribution_chart(results: List[dict]) -> go.Figure:
    """Create score distribution visualization"""
    scores = [r["score"]["overall"] * 100 for r in results]
    fig = go.Figure()
    fig.add_trace(go.Histogram(x=scores, nbinsx=20, marker_color='#4ECDC4', name='Score Distribution'))
    fig.update_layout(
        title="📊 Overall Score Distribution",
        xaxis_title="Score (%)",
        yaxis_title="Number of Candidates",
        height=400,
        showlegend=False
    )
    return fig


def create_experience_vs_score(results: List[dict]) -> go.Figure:
    """Scatter plot of experience vs score"""
    years = [r["years_experience"] for r in results]
    scores = [r["score"]["overall"] * 100 for r in results]
    names = [r["contacts"].get("name", "Unknown") for r in results]
    
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=years, y=scores, mode='markers',
        marker=dict(size=12, color=scores, colorscale='Viridis', showscale=True),
        text=names, hovertemplate='%{text}<br>Years: %{x}<br>Score: %{y}%'
    ))
    fig.update_layout(
        title="📈 Experience vs Performance Score",
        xaxis_title="Years of Experience",
        yaxis_title="Overall Score (%)",
        height=400
    )
    return fig


def create_skill_category_radar(results: List[dict]) -> go.Figure:
    """Aggregate skill category radar for all candidates"""
    categories = ["programming", "data", "ml_ai", "analytics", "cloud", "web", "soft"]
    avg_counts = []
    
    for cat in categories:
        total = sum(len(r["tags"].get(cat, [])) for r in results)
        avg_counts.append(total / len(results) if results else 0)
    
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=avg_counts,
        theta=[c.upper().replace("_", " ") for c in categories],
        fill='toself',
        name='Average Skills',
        line=dict(color='#FF6B6B', width=2)
    ))
    fig.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, max(avg_counts) * 1.2])),
        title="🎯 Average Skill Coverage Across Candidates",
        height=400
    )
    return fig


def create_timeline_gantt(results: List[dict]) -> go.Figure:
    """Create experience timeline Gantt chart for top candidates"""
    gantt_data = []
    colors_palette = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7']
    
    for idx, r in enumerate(results[:5]):  # Top 5
        name = r["contacts"].get("name", f"Candidate {idx+1}")
        for exp in r["experience"]:
            period_str = exp.get("period", "")
            match = YEAR_RANGE_RE.search(period_str)
            if match:
                start_year = int(match.group(1))
                end_str = match.group(2)
                end_year = datetime.now().year if end_str.lower() in ("present", "current") else int(end_str)
                
                gantt_data.append(dict(
                    Task=name,
                    Start=f"{start_year}-01-01",
                    Finish=f"{end_year}-12-31",
                    Resource=exp.get("raw", "Position")[:30]
                ))
    
    if not gantt_data:
        return go.Figure()
    
    df_gantt = pd.DataFrame(gantt_data)
    fig = px.timeline(df_gantt, x_start="Start", x_end="Finish", y="Task", color="Task",
                      hover_data=["Resource"], title="📅 Top 5 Candidates Experience Timeline")
    fig.update_yaxes(categoryorder="total ascending")
    fig.update_layout(height=400, showlegend=False)
    return fig


def create_top_skills_bar(results: List[dict]) -> go.Figure:
    """Bar chart of top skills across all candidates"""
    skills_count = {}
    for r in results:
        for cat, skills in r["tags"].items():
            for skill in skills:
                skills_count[skill] = skills_count.get(skill, 0) + 1
    
    top_15 = sorted(skills_count.items(), key=lambda x: x[1], reverse=True)[:15]
    
    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=[s[1] for s in top_15],
        y=[s[0] for s in top_15],
        orientation='h',
        marker_color='#4ECDC4',
        text=[s[1] for s in top_15],
        textposition='auto'
    ))
    fig.update_layout(
        title="🔥 Top 15 Skills in Candidate Pool",
        xaxis_title="Number of Candidates",
        yaxis_title="Skill",
        height=500
    )
    return fig


def create_comparative_table(results: List[dict]) -> pd.DataFrame:
    """Create side-by-side comparison table for top candidates"""
    top_5 = results[:5]
    comparison = []
    
    for r in top_5:
        comparison.append({
            "Name": r["contacts"].get("name", "Unknown"),
            "Overall %": round(r["score"]["overall"] * 100, 1),
            "JD Fit %": round(r["score"]["jd_fit"] * 100, 1),
            "Years Exp": round(r["years_experience"], 1),
            "Hard Skills": round(r["score"]["breakdown"]["hard_skills"] * 100, 1),
            "Soft Skills": round(r["score"]["breakdown"]["soft_skills"] * 100, 1),
            "Digital": round(r["score"]["breakdown"]["digital_presence"] * 100, 1),
            "LinkedIn": "✅" if r["social"].get("linkedin") else "❌",
            "GitHub": "✅" if r["social"].get("github") else "❌",
            "Gaps": "⚠️" if r.get("gaps") else "✅"
        })
    
    return pd.DataFrame(comparison)


def create_skill_gap_matrix(jd_text: str, results: List[dict]) -> pd.DataFrame:
    """Identify skill gaps between JD and candidates"""
    jd_lower = jd_text.lower()
    
    # Extract key skills from JD
    skill_keywords = []
    for cat, skills in SKILLS.items():
        for skill in skills:
            if skill in jd_lower:
                skill_keywords.append(skill)
    
    gap_matrix = []
    for r in results[:10]:  # Top 10
        cv_text = r["text"].lower()
        name = r["contacts"].get("name", "Unknown")
        
        row = {"Candidate": name}
        for skill in skill_keywords[:15]:  # Top 15 JD skills
            row[skill.title()] = "✅" if skill in cv_text else "❌"
        
        gap_matrix.append(row)
    
    return pd.DataFrame(gap_matrix)


# ========================= Modern Advanced Features =========================
def extract_entities_with_spacy(text: str) -> Dict[str, List[str]]:
    """Extract named entities using spaCy NER"""
    if not SPACY_AVAILABLE or nlp is None:
        return {"persons": [], "orgs": [], "locations": [], "dates": [], "skills": []}
    
    try:
        doc = nlp(text[:1000000])  # Process first 1M chars
        entities = {
            "persons": [ent.text for ent in doc.ents if ent.label_ == "PERSON"],
            "orgs": [ent.text for ent in doc.ents if ent.label_ == "ORG"],
            "locations": [ent.text for ent in doc.ents if ent.label_ in ("GPE", "LOC")],
            "dates": [ent.text for ent in doc.ents if ent.label_ == "DATE"],
            "skills": []
        }
        
        # Extract technical skills from noun chunks
        for chunk in doc.noun_chunks:
            if any(tech in chunk.text.lower() for tech in ["python", "java", "ml", "ai", "cloud", "data"]):
                entities["skills"].append(chunk.text)
        
        return entities
    except Exception:
        return {"persons": [], "orgs": [], "locations": [], "dates": [], "skills": []}


def analyze_text_readability(text: str) -> Dict[str, float]:
    """Analyze CV readability using textstat - ENHANCED with more metrics"""
    if not TEXTSTAT_AVAILABLE:
        return {"flesch_reading_ease": 0, "flesch_kincaid_grade": 0, "reading_time": 0, "word_count": 0, "sentence_count": 0}
    
    try:
        flesch = textstat.flesch_reading_ease(text)
        grade = textstat.flesch_kincaid_grade(text)
        smog = textstat.smog_index(text)
        coleman = textstat.coleman_liau_index(text)
        reading_time = textstat.reading_time(text, ms_per_char=14.69)
        sentences = textstat.sentence_count(text)
        words = textstat.lexicon_count(text, removepunct=True)
        syllables = textstat.syllable_count(text)
        
        # Additional advanced metrics
        automated_readability = textstat.automated_readability_index(text)
        dale_chall = textstat.dale_chall_readability_score(text)
        difficult_words = textstat.difficult_words(text)
        linsear_write = textstat.linsear_write_formula(text)
        gunning_fog = textstat.gunning_fog(text)
        
        return {
            "flesch_reading_ease": flesch,
            "flesch_kincaid_grade": grade,
            "smog_index": smog,
            "coleman_liau": coleman,
            "reading_time": reading_time,
            "sentence_count": sentences,
            "word_count": words,
            "syllable_count": syllables,
            "automated_readability_index": automated_readability,
            "dale_chall_score": dale_chall,
            "difficult_words": difficult_words,
            "linsear_write": linsear_write,
            "gunning_fog": gunning_fog,
            "avg_words_per_sentence": words / sentences if sentences > 0 else 0,
            "avg_syllables_per_word": syllables / words if words > 0 else 0
        }
    except Exception as e:
        return {"flesch_reading_ease": 0, "flesch_kincaid_grade": 0, "reading_time": 0, "word_count": 0, "sentence_count": 0, "error": str(e)}


def detect_language(text: str) -> str:
    """Detect document language"""
    if not LANGDETECT_AVAILABLE:
        return "en"
    
    try:
        return detect(text[:1000])
    except Exception:
        return "en"


def advanced_similarity_with_faiss(jd: str, cvs: List[str]) -> List[float]:
    """Use FAISS for fast similarity search with better accuracy"""
    if not FAISS_AVAILABLE:
        # Fallback to cosine similarity
        embeddings = embed_texts([jd] + cvs)
        jd_emb = embeddings[0]
        cv_embs = embeddings[1:]
        return [cosine(jd_emb, cv_emb) for cv_emb in cv_embs]
    
    try:
        # Embed all texts
        embeddings = embed_texts([jd] + cvs)
        dimension = embeddings.shape[1]
        
        # Create FAISS index
        index = faiss.IndexFlatIP(dimension)  # Inner product (cosine similarity with normalized vectors)
        
        # Normalize vectors
        faiss.normalize_L2(embeddings)
        
        # Add CV embeddings to index
        index.add(embeddings[1:])
        
        # Search for JD against all CVs
        distances, indices = index.search(embeddings[0:1], len(cvs))
        
        return distances[0].tolist()
    except Exception:
        # Fallback
        embeddings = embed_texts([jd] + cvs)
        jd_emb = embeddings[0]
        cv_embs = embeddings[1:]
        return [cosine(jd_emb, cv_emb) for cv_emb in cv_embs]


def extract_certifications(text: str) -> List[str]:
    """DYNAMIC certification extraction - extracts ANY certification from sections + pattern matching"""
    certifications = []
    
    # Strategy 1: Section-based extraction (MOST IMPORTANT - extracts EVERYTHING from cert sections)
    cert_section_markers = [
        r"^(?:professional\s+)?certifications?\s*:?\s*$",
        r"^licenses?\s+(?:and|&|,)\s+certifications?\s*:?\s*$",
        r"^certificates?\s*:?\s*$",
        r"^professional\s+development\s*:?\s*$",
        r"^training\s+(?:and|&)\s+certifications?\s*:?\s*$",
        r"^credentials?\s*:?\s*$",
    ]
    
    lines = text.split('\n')
    in_cert_section = False
    cert_start_line = -1
    
    # Section terminators
    other_section_keywords = ['experience', 'education', 'skills', 'work history', 'employment', 
                              'projects', 'summary', 'profile', 'objective', 'achievements', 
                              'publications', 'references', 'languages', 'interests', 'hobbies']
    
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        
        # Check if entering certification section
        if not in_cert_section:
            for marker in cert_section_markers:
                if re.match(marker, line_stripped, re.I):
                    in_cert_section = True
                    cert_start_line = i
                    break
            continue
        
        # Check if leaving certification section
        if in_cert_section:
            # End on new major section
            is_new_section = False
            for keyword in other_section_keywords:
                if re.match(rf'^{keyword}', line_stripped, re.I):
                    is_new_section = True
                    break
            
            if is_new_section:
                in_cert_section = False
                continue
            
            # End if we see another section header pattern (all caps, followed by colon)
            if re.match(r'^[A-Z\s]{3,30}:?$', line_stripped) and len(line_stripped) < 40:
                in_cert_section = False
                continue
            
            # Extract certification from this line (DYNAMIC - takes anything)
            if line_stripped and len(line_stripped) > 2:
                # Skip pure dates or page numbers
                if re.match(r'^\d{1,3}$', line_stripped):  # Page number
                    continue
                if re.match(r'^\d{4}\s*[-–]\s*\d{4}$', line_stripped):  # Just dates
                    continue
                if re.match(r'^\d{1,2}/\d{4}$', line_stripped):  # Date format
                    continue
                
                # Clean the line
                cert_line = line_stripped
                
                # Remove common bullets and markers
                cert_line = re.sub(r'^[\u2022\u25cf\u25e6\u25aa\u25ab\-\*•◦▪▫→➢⮚]\s*', '', cert_line)
                cert_line = re.sub(r'^\d+[.)\s]+', '', cert_line)  # Remove numbering
                
                # Remove trailing dates in common formats
                cert_line = re.sub(r',?\s*\d{4}\s*$', '', cert_line)
                cert_line = re.sub(r',?\s*\d{1,2}/\d{4}\s*$', '', cert_line)
                cert_line = re.sub(r',?\s*\(\d{4}\)\s*$', '', cert_line)
                cert_line = re.sub(r'\s*[-–]\s*\d{4}\s*$', '', cert_line)
                
                # Remove "Issued by" or "by" suffixes
                cert_line = re.sub(r'\s*[-–]\s*issued\s+by\s+[^,]+$', '', cert_line, flags=re.I)
                cert_line = re.sub(r'\s*\(issued\s+by\s+[^)]+\)\s*$', '', cert_line, flags=re.I)
                
                cert_line = cert_line.strip()
                
                # Must be reasonable length
                if 3 <= len(cert_line) <= 200:
                    # Must have at least one letter
                    if re.search(r'[a-zA-Z]', cert_line):
                        certifications.append(cert_line)
    
    # Strategy 2: Keyword-based pattern extraction (backup for certs mentioned outside sections)
    cert_keywords = [
        r"certified", r"certification", r"certificate", r"credential", r"licensed", 
        r"accredited", r"qualified", r"professional\s+\w+"
    ]
    
    # Look for lines containing certification keywords
    for line in lines:
        line_stripped = line.strip()
        if 10 <= len(line_stripped) <= 200:
            # Check if contains cert keyword
            has_cert_keyword = any(re.search(rf'\b{kw}\b', line_stripped, re.I) for kw in cert_keywords)
            
            if has_cert_keyword:
                # Extract potential certification
                cert_candidate = line_stripped
                
                # Clean up
                cert_candidate = re.sub(r'^[\u2022\u25cf\u25e6\u25aa\u25ab\-\*•◦▪▫→]\s*', '', cert_candidate)
                cert_candidate = re.sub(r'^\d+[.)\s]+', '', cert_candidate)
                cert_candidate = re.sub(r',?\s*\d{4}\s*$', '', cert_candidate)
                cert_candidate = cert_candidate.strip()
                
                if 5 <= len(cert_candidate) <= 200:
                    certifications.append(cert_candidate)
    
    # Strategy 3: Common certification acronyms (even if standalone)
    acronym_patterns = [
        r'\b(AWS|Azure|GCP|PMP|CAPM|CSM|CSPO|PSM|CISSP|CISM|CISA|CEH|OSCP|' +
        r'CCNA|CCNP|CCIE|RHCSA|RHCE|CKA|CKAD|CKS|CompTIA|ITIL|' +
        r'OCA|OCP|CPA|CFA|CFP|CMA|FRM)(?:\s+Certified)?(?:\s+[A-Za-z\s]+)?'
    ]
    
    for pattern in acronym_patterns:
        matches = re.findall(pattern, text, re.I)
        for match in matches:
            if 2 <= len(match) <= 100:
                certifications.append(match.strip())
    
    # Strategy 4: "Certified X" or "X Certification" patterns
    generic_patterns = [
        r'Certified\s+[A-Z][A-Za-z\s]{2,60}(?:Engineer|Developer|Administrator|Architect|Specialist|Professional|Practitioner|Consultant|Analyst|Expert)',
        r'[A-Z][A-Za-z\s]{2,40}\s+Certification',
        r'[A-Z][A-Za-z\s]{2,40}\s+Professional\s+Certificate',
    ]
    
    for pattern in generic_patterns:
        matches = re.findall(pattern, text, re.I)
        for match in matches:
            if 5 <= len(match) <= 150:
                certifications.append(match.strip())
    
    # Deduplicate and clean
    unique_certs = []
    seen_lower = set()
    
    for cert in certifications:
        cert = cert.strip()
        cert_lower = cert.lower()
        
        # Skip if empty or too short
        if len(cert) < 2:
            continue
        
        # Skip if already seen
        if cert_lower in seen_lower:
            continue
        
        # Skip if it's just common words
        skip_words = {'certified', 'certification', 'certificate', 'professional', 'licensed', 'the', 'and', 'or'}
        if cert_lower in skip_words:
            continue
        
        unique_certs.append(cert)
        seen_lower.add(cert_lower)
    
    return unique_certs


def calculate_cv_quality_score(text: str, readability: Dict, entities: Dict) -> float:
    """Calculate overall CV quality based on multiple factors"""
    score = 0.0
    
    # Length check (800-3000 words is ideal)
    word_count = len(text.split())
    if 800 <= word_count <= 3000:
        score += 0.2
    elif 500 <= word_count < 800 or 3000 < word_count <= 4000:
        score += 0.1
    
    # Readability (Flesch score 60-80 is ideal)
    flesch = readability.get("flesch_score", 0)
    if 60 <= flesch <= 80:
        score += 0.2
    elif 40 <= flesch < 60 or 80 < flesch <= 90:
        score += 0.1
    
    # Structure (presence of sections)
    has_experience = bool(re.search(r"experience|work history", text, re.I))
    has_education = bool(re.search(r"education|academic", text, re.I))
    has_skills = bool(re.search(r"skills|competencies", text, re.I))
    
    score += 0.15 if has_experience else 0
    score += 0.15 if has_education else 0
    score += 0.1 if has_skills else 0
    
    # Entity richness
    if entities.get("orgs") and len(entities["orgs"]) >= 2:
        score += 0.1
    if entities.get("locations"):
        score += 0.05
    
    # Contact info
    has_email = bool(EMAIL_RE.search(text))
    has_phone = bool(PHONE_RE.search(text))
    score += 0.05 if has_email else 0
    score += 0.05 if has_phone else 0
    
    return min(score, 1.0)


def create_enhanced_radar_chart(results: List[dict]) -> go.Figure:
    """Create enhanced multi-dimensional radar chart with more metrics"""
    top_5 = results[:5]
    
    fig = go.Figure()
    colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7']
    
    categories = ['JD Fit', 'Hard Skills', 'Experience', 'Soft Skills', 'Digital', 'Formatting', 'Quality']
    
    for idx, r in enumerate(top_5):
        name = r["contacts"].get("name", f"Candidate {idx+1}")
        breakdown = r["score"]["breakdown"]
        
        values = [
            r["score"]["jd_fit"] * 100,
            breakdown["hard_skills"] * 100,
            breakdown["experience"] * 100,
            breakdown["soft_skills"] * 100,
            breakdown["digital_presence"] * 100,
            breakdown["formatting"] * 100,
            r.get("cv_quality_score", 0.7) * 100
        ]
        
        fig.add_trace(go.Scatterpolar(
            r=values,
            theta=categories,
            fill='toself',
            name=name,
            line=dict(color=colors[idx % len(colors)], width=2)
        ))
    
    fig.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
        title="🎯 Multi-Dimensional Candidate Analysis",
        height=500,
        showlegend=True
    )
    return fig


# ========================= Classic Visuals Utilities =========================
def extract_text_from_pdf(file) -> str:
    if _PdfReader is None:
        return ""
    pdf = _PdfReader(file)
    text = ""
    for page in pdf.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text
    return text


def rank_resumes_tfidf(job_description, resumes):
    documents = [job_description] + resumes
    vectorizer = TfidfVectorizer().fit_transform(documents)
    vectors = vectorizer.toarray()
    job_description_vector = vectors[0]
    resume_vectors = vectors[1:]
    cosine_similarities = cosine_similarity([job_description_vector], resume_vectors).flatten()
    scores = (cosine_similarities * 100).round(2)
    return scores


def create_parallel_coordinates(results, resumes, job_description):
    parallel_data = []
    skill_categories = {
        'Technical': ['python', 'java', 'sql', 'programming', 'coding', 'software'],
        'Experience': ['years', 'experience', 'senior', 'lead', 'manager', 'director'],
        'Education': ['degree', 'university', 'college', 'bachelor', 'master', 'phd'],
        'Communication': ['communication', 'presentation', 'writing', 'team', 'collaboration']
    }
    for idx, (_, candidate) in enumerate(results.iterrows()):
        resume_text = resumes[idx].lower()
        dimensions = {'VitaSort_Score': candidate['🎯 VitaSort Score'], 'Ranking': candidate['📈 Ranking']}
        for category, keywords in skill_categories.items():
            score = sum(resume_text.count(keyword) for keyword in keywords)
            dimensions[f'{category}_Score'] = min(score * 10, 100)
        dimensions['Resume'] = candidate['🍎 Resume'].replace('.pdf', '')
        parallel_data.append(dimensions)
    parallel_df = pd.DataFrame(parallel_data)
    fig = go.Figure(data=go.Parcoords(
        line=dict(color=parallel_df['VitaSort_Score'], colorscale='viridis', showscale=True,
                  cmin=parallel_df['VitaSort_Score'].min(), cmax=parallel_df['VitaSort_Score'].max()),
        dimensions=list([
            dict(range=[0, 100], constraintrange=[0, 100], label="VitaSort Score", values=parallel_df['VitaSort_Score']),
            dict(range=[0, 100], label="Technical Skills", values=parallel_df['Technical_Score']),
            dict(range=[0, 100], label="Experience Level", values=parallel_df['Experience_Score']),
            dict(range=[0, 100], label="Education", values=parallel_df['Education_Score']),
            dict(range=[0, 100], label="Communication", values=parallel_df['Communication_Score']),
            dict(range=[1, len(results)], label="Ranking", values=parallel_df['Ranking'], tickvals=list(range(1, len(results)+1)))
        ])
    ))
    fig.update_layout(title="📊 Parallel Coordinates Analysis", height=600)
    return fig


def create_radar_chart(results, resumes, job_description):
    skill_categories = {
        'Technical Skills': ['python', 'java', 'sql', 'programming', 'coding', 'software'],
        'Data Science': ['machine learning', 'data science', 'analytics', 'statistics', 'ai'],
        'Business Skills': ['management', 'leadership', 'strategy', 'business', 'project'],
        'Communication': ['communication', 'presentation', 'writing', 'collaboration', 'team'],
        'Experience': ['years', 'experience', 'senior', 'lead', 'manager', 'director']
    }
    top_candidates = results.head(3)
    fig = go.Figure()
    colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7']
    for idx, (_, candidate) in enumerate(top_candidates.iterrows()):
        resume_idx = results[results['🍎 Resume'] == candidate['🍎 Resume']].index[0]
        resume_text = resumes[resume_idx].lower()
        category_scores = []
        for _, keywords in skill_categories.items():
            score = sum(resume_text.count(keyword) for keyword in keywords)
            category_scores.append(min(score * 10, 100))
        fig.add_trace(go.Scatterpolar(r=category_scores, theta=list(skill_categories.keys()), fill='toself',
                                      name=f"{candidate['🍎 Resume'].replace('.pdf', '')} (Score: {candidate['🎯 VitaSort Score']})",
                                      line=dict(color=colors[idx % len(colors)], width=3)))
    fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 100])), height=600, showlegend=True)
    return fig


def create_word_cloud_comparison(job_description, top_resume_text):
    def clean_text(text):
        t = re.sub(r'[^\w\s]', '', text.lower())
        return ' '.join([w for w in t.split() if len(w) > 2])
    job_text_clean = clean_text(job_description)
    resume_text_clean = clean_text(top_resume_text)
    job_wordcloud = WordCloud(width=400, height=300, background_color='white', colormap='plasma', max_words=100, min_font_size=10).generate(job_text_clean)
    resume_wordcloud = WordCloud(width=400, height=300, background_color='white', colormap='viridis', max_words=100, min_font_size=10).generate(resume_text_clean)
    return job_wordcloud, resume_wordcloud


def create_skills_heatmap(job_description, resumes, filenames):
    skills = ['python', 'java', 'sql', 'machine learning', 'data science', 'analytics', 'project management', 'communication', 'leadership', 'problem solving', 'teamwork', 'excel', 'powerbi', 'tableau']
    skill_matrix = []
    labels = ['Job Description'] + [f.replace('.pdf', '') for f in filenames]
    all_texts = [job_description.lower()] + [r.lower() for r in resumes]
    for text in all_texts:
        scores = []
        for skill in skills:
            count = text.count(skill.lower())
            scores.append(min(count * 2, 10))
        skill_matrix.append(scores)
    fig = px.imshow(skill_matrix, labels=dict(x="Skills", y="Documents", color="Match Score"),
                    x=[s.title() for s in skills], y=labels, color_continuous_scale="Turbo",
                    title="🎯 Skills Matching Heatmap")
    fig.update_layout(height=600)
    return fig


# ========================= Practical Features =========================
def get_status_badge(overall_score: float, filename: str) -> str:
    """Generate status badge based on score and saved status"""
    if filename in st.session_state.candidate_status:
        status = st.session_state.candidate_status[filename]
        badge_map = {
            "shortlisted": "🌟 Shortlisted",
            "interview": "📞 Interview",
            "rejected": "❌ Rejected",
            "offer": "💼 Offer",
            "hired": "✅ Hired"
        }
        return badge_map.get(status, "📋 New")
    
    # Auto-status based on score
    if overall_score >= 0.80:
        return "🟢 Strong"
    elif overall_score >= 0.60:
        return "🟡 Average"
    elif overall_score >= 0.40:
        return "🟠 Below Avg"
    else:
        return "🔴 Weak"


def add_candidate_note(filename: str, note: str):
    """Add note to candidate"""
    if filename not in st.session_state.candidate_notes:
        st.session_state.candidate_notes[filename] = []
    st.session_state.candidate_notes[filename].append({
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "note": note
    })


def toggle_bookmark(filename: str):
    """Toggle bookmark status"""
    if filename in st.session_state.bookmarked:
        st.session_state.bookmarked.remove(filename)
    else:
        st.session_state.bookmarked.add(filename)


def update_candidate_status(filename: str, new_status: str):
    """Update candidate workflow status"""
    st.session_state.candidate_status[filename] = new_status


def generate_quick_stats(results: List[Dict]) -> str:
    """Generate quick stats for live HUD"""
    total = len(results)
    strong = len([r for r in results if r["score"]["overall"] >= 0.80])
    avg_score = sum([r["score"]["overall"] for r in results]) / total if total > 0 else 0
    bookmarked_count = len([r for r in results if r["filename"] in st.session_state.bookmarked])
    return f"🔍 {total} candidates | 🟢 {strong} strong | ⭐ {bookmarked_count} bookmarked | 📊 {avg_score:.0%} avg"


# ========================= Session State Init =========================
if 'candidate_status' not in st.session_state:
    st.session_state.candidate_status = {}
if 'candidate_notes' not in st.session_state:
    st.session_state.candidate_notes = {}
if 'bookmarked' not in st.session_state:
    st.session_state.bookmarked = set()
if 'score_threshold' not in st.session_state:
    st.session_state.score_threshold = 0.0
if 'custom_weights' not in st.session_state:
    st.session_state.custom_weights = {"hard": 0.40, "exp": 0.30, "soft": 0.10, "digital": 0.10, "format": 0.10}
if 'filter_presets' not in st.session_state:
    st.session_state.filter_presets = {
        "Senior Dev": {"min_years": 5.0, "threshold": 0.70, "skills": "python senior lead", "location": "", "boolean": ""},
        "Junior Dev": {"min_years": 0.0, "threshold": 0.50, "skills": "junior entry", "location": "", "boolean": ""},
        "Manager": {"min_years": 3.0, "threshold": 0.65, "skills": "management leadership team", "location": "", "boolean": ""}
    }
if 'preset_templates' not in st.session_state:
    st.session_state.preset_templates = {
        "🐍 Python Developer": {"min_years": 2.0, "threshold": 0.65, "skills": "python django flask fastapi", "boolean": "python AND (django OR flask OR fastapi)", "location": ""},
        "⚛️ React Developer": {"min_years": 2.0, "threshold": 0.65, "skills": "react javascript typescript", "boolean": "react AND javascript", "location": ""},
        "🤖 ML Engineer": {"min_years": 3.0, "threshold": 0.75, "skills": "machine learning tensorflow pytorch", "boolean": "(machine learning OR ML) AND (tensorflow OR pytorch)", "location": ""},
        "☁️ Cloud Architect": {"min_years": 5.0, "threshold": 0.80, "skills": "aws azure gcp kubernetes", "boolean": "(aws OR azure OR gcp) AND (cloud OR kubernetes)", "location": ""},
        "📊 Data Scientist": {"min_years": 3.0, "threshold": 0.70, "skills": "python data science statistics", "boolean": "data science AND (python OR R) AND statistics", "location": ""},
        "🔒 Security Engineer": {"min_years": 4.0, "threshold": 0.75, "skills": "security penetration testing", "boolean": "security AND (penetration OR infosec)", "location": ""},
        "📱 Mobile Developer": {"min_years": 2.0, "threshold": 0.65, "skills": "android ios react native", "boolean": "(android OR ios OR react native)", "location": ""},
        "🎨 UI/UX Designer": {"min_years": 2.0, "threshold": 0.60, "skills": "figma sketch design ui ux", "boolean": "(ui OR ux OR design) AND (figma OR sketch)", "location": ""},
        "🚀 DevOps Engineer": {"min_years": 3.0, "threshold": 0.70, "skills": "devops ci/cd docker kubernetes", "boolean": "devops AND (docker OR kubernetes) AND ci/cd", "location": ""},
        "💼 Product Manager": {"min_years": 4.0, "threshold": 0.70, "skills": "product management agile roadmap", "boolean": "product AND management AND (agile OR scrum)", "location": ""}
    }
if 'saved_filter_presets' not in st.session_state:
    st.session_state.saved_filter_presets = {}  # User-created presets
if 'last_results' not in st.session_state:
    st.session_state.last_results = []
if 'live_stats' not in st.session_state:
    st.session_state.live_stats = {}
if 'quick_search' not in st.session_state:
    st.session_state.quick_search = ""
if 'view_mode' not in st.session_state:
    st.session_state.view_mode = "all"

# ========================= SIDEBAR =========================
st.sidebar.markdown("---")
st.sidebar.markdown("Developer")
st.sidebar.markdown("**LABIB BIN SHAHED**")
st.sidebar.markdown("labib-x@protonmail.com")
st.sidebar.markdown("---")

# ========================= MAIN CONTENT AREA =========================
st.markdown("---")

# ========================= CONTROL PANEL =========================
control_tab1, control_tab2, control_tab3, control_tab4, control_tab5, control_tab6 = st.tabs([
    "🎛️ Mode", "📊 Analytics", "⚡ Presets", "⚙️ Scoring", "🔬 Features", "⚡ Actions"
])

# ====== Tab 1: Mode Selection ======
with control_tab1:
    st.markdown("### 🎛️ Mode Selection")
    st.info("Choose between **Forensics Mode** (detailed live analysis) or **Default Mode** (simplified view).")
    
    mode = st.radio(
        "Select analysis mode",
        ["🔬 Forensics Mode", "📊 Default Mode"],
        horizontal=True,
        key="mode_selector",
        label_visibility="collapsed"
    )
    # Clean mode name
    mode = "Forensics" if "Forensics" in mode else "Default"
    
    # Mode descriptions
    if "Forensics" in mode:
        st.success("""
        **🔬 Forensics Mode Active**
        - Live HUD with real-time metrics
        - Comprehensive analytics dashboard
        - Advanced filtering and search
        - Detailed candidate comparisons
        """)
    else:
        st.success("""
        **📊 Default Mode Active**
        - Simplified interface
        - Quick candidate review
        - Essential metrics only
        - Streamlined workflow
        """)

if mode == "Forensics":
    # Live HUD Header
    st.markdown('<div style="position: sticky; top: 0; z-index: 999; background: white; padding: 10px 0; border-bottom: 2px solid #f0f0f0;">', unsafe_allow_html=True)
    if st.session_state.last_results:
        hud_cols = st.columns([2, 1, 1, 1, 1])
        with hud_cols[0]:
            st.markdown(f"### 🎯 VitaSort Live HUD")
        with hud_cols[1]:
            total = len(st.session_state.last_results)
            st.metric("📊 Total", total, delta=None, delta_color="off")
        with hud_cols[2]:
            strong = len([r for r in st.session_state.last_results if r["score"]["overall"] >= 0.80])
            st.metric("🟢 Strong", strong, delta=f"{(strong/total*100):.0f}%" if total > 0 else "0%")
        with hud_cols[3]:
            bookmarked = len([r for r in st.session_state.last_results if r["filename"] in st.session_state.bookmarked])
            st.metric("⭐ Starred", bookmarked, delta=f"{(bookmarked/total*100):.0f}%" if total > 0 else "0%")
        with hud_cols[4]:
            avg = sum([r["score"]["overall"] for r in st.session_state.last_results]) / total if total > 0 else 0
            st.metric("📈 Avg", f"{avg:.0%}", delta="Good" if avg >= 0.70 else "Low")
    st.markdown('</div>', unsafe_allow_html=True)

    # ====== Tab 2: Live Pool Analytics ======
    with control_tab2:
        # Live Pool Analytics
        if st.session_state.last_results:
            results = st.session_state.last_results
            st.markdown("### 📊 Live Pool Stats")
            pool_col1, pool_col2, pool_col3, pool_col4 = st.columns(4)
            with pool_col1:
                st.metric("Total Candidates", len(results))
            with pool_col2:
                strong_count = len([r for r in results if r["score"]["overall"] >= 0.80])
                st.metric("🟢 Strong (≥80%)", strong_count)
            with pool_col3:
                bookmarked_count = len([r for r in results if r["filename"] in st.session_state.bookmarked])
                st.metric("⭐ Starred", bookmarked_count)
            with pool_col4:
                avg_score = sum([r["score"]["overall"] for r in results]) / len(results)
                st.metric("Avg Score", f"{avg_score:.0%}")
            
            # Status Distribution
            st.markdown("### 📈 Status Distribution")
            status_counts = {"new": 0, "shortlisted": 0, "interview": 0, "rejected": 0, "offer": 0}
            for r in results:
                status = st.session_state.candidate_status.get(r["filename"], "new")
                if status in status_counts:
                    status_counts[status] += 1
            
            status_cols = st.columns(5)
            for idx, (status, count) in enumerate(status_counts.items()):
                with status_cols[idx]:
                    pct = (count / len(results)) * 100 if len(results) > 0 else 0
                    st.metric(status.title(), count, delta=f"{pct:.0f}%")
            
            st.markdown("---")
        
        # Quick Search & View Mode
        search_col1, search_col2 = st.columns(2)
        with search_col1:
            st.markdown("### 🔍 Quick Search")
            quick_search = st.text_input("Search candidates...", key="quick_search_input", placeholder="Name, email, skills...")
            if quick_search:
                st.session_state.quick_search = quick_search
        
        with search_col2:
            st.markdown("### 👁️ View Mode")
            view_options = ["All", "⭐ Bookmarked", "🟢 Strong (≥80%)", "🟡 Average (60-80%)", "🔴 Weak (<60%)", "🌟 Shortlisted", "📞 Interview", "❌ Rejected"]
            st.session_state.view_mode = st.selectbox("Filter View", view_options, key="view_mode_select")
    
    # ====== Tab 3: Filter Presets ======
    with control_tab3:
        st.markdown("### ⚡ Filter Presets Management")
        
        # Quick Presets
        st.markdown("**Quick Apply:**")
        qp_col1, qp_col2, qp_col3, qp_col4 = st.columns(4)
        with qp_col1:
            if st.button("👔 Senior Dev", use_container_width=True):
                preset = st.session_state.filter_presets["Senior Dev"]
                st.session_state.score_threshold = preset["threshold"]
                st.rerun()
        with qp_col2:
            if st.button("🎓 Junior Dev", use_container_width=True):
                preset = st.session_state.filter_presets["Junior Dev"]
                st.session_state.score_threshold = preset["threshold"]
                st.rerun()
        with qp_col3:
            if st.button("💼 Manager", use_container_width=True):
                preset = st.session_state.filter_presets["Manager"]
                st.session_state.score_threshold = preset["threshold"]
                st.rerun()
        with qp_col4:
            if st.button("🔄 Clear All", use_container_width=True):
                st.session_state.score_threshold = 0.0
                st.session_state.quick_search = ""
                st.rerun()
        
        st.markdown("---")
        
        # Template Library
        preset_col1, preset_col2 = st.columns([2, 1])
        
        with preset_col1:
            st.markdown("### 📚 Template Library")
            template_select = st.selectbox(
                "Choose job role template",
                options=["Select..."] + list(st.session_state.preset_templates.keys()),
                key="template_select"
            )
            
            if template_select != "Select...":
                template = st.session_state.preset_templates[template_select]
                st.info(f"**Preview:** {template['min_years']}+ yrs | {template['threshold']:.0%} threshold | Skills: {template['skills']}")
            
            temp_col1, temp_col2 = st.columns(2)
            with temp_col1:
                if st.button("✅ Apply Template", use_container_width=True, disabled=(template_select == "Select...")):
                    if template_select != "Select...":
                        template = st.session_state.preset_templates[template_select]
                        st.session_state.score_threshold = template["threshold"]
                        if 'yrs' in st.session_state:
                            st.session_state.yrs = template["min_years"]
                        if 'ftq' in st.session_state:
                            st.session_state.ftq = template["skills"]
                        if 'bq' in st.session_state:
                            st.session_state.bq = template["boolean"]
                        st.success(f"✅ Applied: {template_select}")
                        st.rerun()
            with temp_col2:
                if st.button("💾 Save to My Presets", use_container_width=True, disabled=(template_select == "Select...")):
                    if template_select != "Select...":
                        st.session_state.saved_filter_presets[template_select] = st.session_state.preset_templates[template_select]
                        st.success(f"Saved!")
                        st.rerun()
        
        with preset_col2:
            st.markdown("### 💾 My Saved Presets")
            if st.session_state.saved_filter_presets:
                selected_preset = st.selectbox(
                    "Load saved preset",
                    options=["Select..."] + list(st.session_state.saved_filter_presets.keys()),
                    key="load_preset_select"
                )
                
                load_col1, load_col2 = st.columns([2, 1])
                with load_col1:
                    if st.button("📂 Load", use_container_width=True, disabled=(selected_preset == "Select...")):
                        if selected_preset != "Select...":
                            preset = st.session_state.saved_filter_presets[selected_preset]
                            st.session_state.score_threshold = preset["threshold"]
                            st.success(f"✅ Loaded: {selected_preset}")
                            st.rerun()
                with load_col2:
                    if st.button("🗑️", use_container_width=True, disabled=(selected_preset == "Select...")):
                        if selected_preset != "Select...":
                            del st.session_state.saved_filter_presets[selected_preset]
                            st.success(f"Deleted")
                            st.rerun()
                
                st.caption(f"📋 {len(st.session_state.saved_filter_presets)} saved preset(s)")
            else:
                st.info("No saved presets yet")
        
        st.markdown("---")
        
        # Create New Preset
        with st.expander("➕ Create New Preset"):
            preset_name = st.text_input("Preset Name", placeholder="e.g., Senior Python Dev", key="new_preset_name")
            
            preset_config_col1, preset_config_col2 = st.columns(2)
            with preset_config_col1:
                preset_min_years = st.number_input("Min Years", 0.0, 50.0, 0.0, 0.5, key="preset_years")
                preset_skills = st.text_input("Skills", placeholder="python, react...", key="preset_skills")
                preset_location = st.text_input("Location", placeholder="City name", key="preset_location")
            with preset_config_col2:
                preset_threshold = st.slider("Threshold %", 0, 100, 50, 5, key="preset_threshold")
                preset_boolean = st.text_input("Boolean Query", placeholder="(python OR java) AND cloud", key="preset_boolean")
            
            if st.button("💾 Save Preset", use_container_width=True, disabled=not preset_name):
                if preset_name:
                    st.session_state.saved_filter_presets[preset_name] = {
                        "min_years": preset_min_years,
                        "threshold": preset_threshold / 100.0,
                        "skills": preset_skills,
                        "boolean": preset_boolean,
                        "location": preset_location
                    }
                    st.success(f"✅ Saved: {preset_name}")
                    st.rerun()
        
        # Import/Export
        with st.expander("📤 Import/Export Presets"):
            ie_col1, ie_col2 = st.columns(2)
            with ie_col1:
                st.markdown("**Export:**")
                if st.button("📤 Export All Presets", use_container_width=True):
                    import json
                    presets_json = json.dumps(st.session_state.saved_filter_presets, indent=2)
                    st.download_button(
                        "⬇️ Download JSON",
                        presets_json,
                        "vitasort_presets.json",
                        "application/json",
                        use_container_width=True
                    )
            with ie_col2:
                st.markdown("**Import:**")
                uploaded_presets = st.file_uploader("Upload JSON", type=["json"], key="import_presets_file")
                if uploaded_presets:
                    try:
                        import json
                        imported_data = json.load(uploaded_presets)
                        if isinstance(imported_data, dict):
                            st.session_state.saved_filter_presets.update(imported_data)
                            st.success(f"✅ Imported {len(imported_data)} preset(s)!")
                            st.rerun()
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
    
    # ====== Tab 4: Scoring Configuration ======
    with control_tab4:
        st.markdown("### ⚙️ Scoring Configuration")
        
        score_col1, score_col2 = st.columns([1, 2])
        with score_col1:
            st.markdown("**Score Threshold:**")
            score_threshold = st.slider("Min Score %", 0, 100, int(st.session_state.score_threshold * 100), 5, key="score_slider")
            st.session_state.score_threshold = score_threshold / 100.0
            st.info(f"Filtering candidates below {score_threshold}%")
        
        with score_col2:
            st.markdown("**Custom Weights:**")
            weight_col1, weight_col2 = st.columns(2)
            with weight_col1:
                st.session_state.custom_weights["hard"] = st.slider("🔧 Hard Skills", 0.0, 1.0, st.session_state.custom_weights["hard"], 0.05, key="hw")
                st.session_state.custom_weights["exp"] = st.slider("📅 Experience", 0.0, 1.0, st.session_state.custom_weights["exp"], 0.05, key="ew")
                st.session_state.custom_weights["soft"] = st.slider("💬 Soft Skills", 0.0, 1.0, st.session_state.custom_weights["soft"], 0.05, key="sw")
            with weight_col2:
                st.session_state.custom_weights["digital"] = st.slider("🌐 Digital Presence", 0.0, 1.0, st.session_state.custom_weights["digital"], 0.05, key="dw")
                st.session_state.custom_weights["format"] = st.slider("📄 Formatting", 0.0, 1.0, st.session_state.custom_weights["format"], 0.05, key="fw")
                total = sum(st.session_state.custom_weights.values())
                if abs(total - 1.0) > 0.05:
                    st.warning(f"⚠️ Total: {total:.2f} (should be 1.0)")
                else:
                    st.success(f"✓ Total: {total:.2f}")
    
    # ====== Tab 5: Advanced Features ======
    with control_tab5:
        st.markdown("### 🔬 Advanced Features")
        
        st.markdown("**🎯 AI-Powered Features:**")
        ai_col1, ai_col2, ai_col3 = st.columns(3)
        with ai_col1:
            enable_ner = st.checkbox("🤖 Named Entity Recognition", value=True, key="ner", help="Extract persons, orgs, locations using spaCy")
            enable_readability = st.checkbox("📖 Readability Analysis", value=True, key="readability", help="Analyze CV quality and reading level")
        with ai_col2:
            enable_lang_detect = st.checkbox("🌍 Language Detection", value=True, key="lang", help="Detect CV language")
            enable_cert_extract = st.checkbox("🎓 Certification Extraction", value=True, key="certs", help="Extract professional certifications")
        with ai_col3:
            enable_faiss = st.checkbox("⚡ FAISS Similarity", value=True, key="faiss", help="Fast similarity search with FAISS")
            enable_quality = st.checkbox("✨ CV Quality Score", value=True, key="quality", help="Calculate overall CV quality")
        
        st.markdown("---")
        st.markdown("**🔍 Classic Features:**")
        feat_col1, feat_col2, feat_col3 = st.columns(3)
        with feat_col1:
            enable_salary_est = st.checkbox("💰 Salary Estimation", value=True, key="sal")
            enable_github = st.checkbox("💻 GitHub Activity", value=True, key="gh")
        with feat_col2:
            enable_hibp = st.checkbox("🔐 HIBP Breach Check", value=False, key="hibp")
            enable_proficiency = st.checkbox("🎯 Skill Proficiency", value=True, key="prof")
        with feat_col3:
            enable_blind = st.checkbox("👤 Blind Screening", value=False, key="blind")
            enable_hopping = st.checkbox("🔄 Job Hopping Detection", value=True, key="hop")
    
    # ====== Tab 6: Quick Actions ======
    with control_tab6:
        st.markdown("### ⚡ Quick Actions")
        
        action_row1_col1, action_row1_col2, action_row1_col3, action_row1_col4 = st.columns(4)
        with action_row1_col1:
            if st.button("🔄 Refresh Data", use_container_width=True):
                st.rerun()
        with action_row1_col2:
            if st.button("🗑️ Clear Session", use_container_width=True):
                st.session_state.candidate_status = {}
                st.session_state.candidate_notes = {}
                st.session_state.bookmarked = set()
                st.success("Session cleared!")
                st.rerun()
        with action_row1_col3:
            if st.button("📥 Export Results", use_container_width=True) and st.session_state.last_results:
                df = pd.DataFrame([{"Filename": r["filename"], "Score": r["score"]["overall"]} for r in st.session_state.last_results])
                st.download_button("⬇️ Download CSV", df.to_csv(index=False), "export.csv", use_container_width=True)
        with action_row1_col4:
            if st.session_state.last_results:
                st.metric("Session", f"{len(st.session_state.last_results)} CVs")
    
    st.markdown("---")
    
    # ========================= JOB DESCRIPTION & UPLOAD =========================
    st.header("Job Description")
    jd_text = st.text_area("Paste Job Description", height=160, key="jd_legendary")
    st.header("Upload CVs (PDF/DOCX/TXT/RTF)")
    files = st.file_uploader("Upload files", type=["pdf", "docx", "txt", "rtf"], accept_multiple_files=True, key="files_legendary")
    
    # Quick Apply Preset Bar
    if st.session_state.saved_filter_presets or st.session_state.filter_presets:

        preset_apply_cols = st.columns([3, 1, 1])
        with preset_apply_cols[0]:
            all_presets = {**st.session_state.filter_presets, **st.session_state.saved_filter_presets}
            selected_apply_preset = st.selectbox(
                "Apply preset to filters below",
                options=["None"] + list(all_presets.keys()),
                key="quick_apply_preset"
            )
        with preset_apply_cols[1]:
            if st.button("✅ Apply", use_container_width=True, disabled=(selected_apply_preset == "None")):
                if selected_apply_preset != "None":
                    preset = all_presets[selected_apply_preset]
                    # Update session state filter values
                    if 'yrs' in st.session_state:
                        st.session_state.yrs = preset.get("min_years", 0.0)
                    if 'ftq' in st.session_state:
                        st.session_state.ftq = preset.get("skills", "")
                    if 'bq' in st.session_state:
                        st.session_state.bq = preset.get("boolean", "")
                    if 'jloc' in st.session_state:
                        st.session_state.jloc = preset.get("location", "")
                    st.session_state.score_threshold = preset.get("threshold", 0.0)
                    st.success(f"✅ Applied preset: {selected_apply_preset}")
                    st.rerun()
        with preset_apply_cols[2]:
            if st.button("📤 Export Presets", use_container_width=True):
                import json
                presets_json = json.dumps(st.session_state.saved_filter_presets, indent=2)
                st.download_button(
                    "⬇️ Download JSON",
                    presets_json,
                    "vitasort_presets.json",
                    "application/json",
                    use_container_width=True
                )
    
    # Live Filters with instant feedback
    st.markdown("### 🔍 Live Filters")
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        ft_query = st.text_input("Full-text filter", key="ftq", placeholder="Search keywords...")
    with c2:
        bool_query = st.text_input("Boolean filter (AND/OR/NOT)", key="bq", placeholder="python AND (senior OR lead)")
    with c3:
        min_years = st.number_input("Min years", min_value=0.0, max_value=50.0, step=0.5, value=0.0, key="yrs")
    with c4:
        jd_location = st.text_input("Job Location (km radius)", key="jloc", placeholder="City name")
    
    if st.button("Analyze", key="run_legendary") and files and jd_text:
        results = []
        
        # Progress tracking
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        total_files = len(files)
        for idx, f in enumerate(files):
            status_text.text(f"Processing {idx+1}/{total_files}: {f.name}")
            progress_bar.progress((idx + 1) / total_files)
            
            data = f.read()
            if f.name.lower().endswith(".pdf"):
                text, meta = pdf_to_text(data)
            elif f.name.lower().endswith(".docx"):
                text, meta = docx_to_text(data), None
            elif f.name.lower().endswith(".rtf"):
                text, meta = rtf_to_text(data), None
            else:
                text, meta = data.decode("utf-8", errors="ignore"), None
            contacts = extract_contacts(text)
            social = extract_social(text)
            edu = extract_education(text)
            exp, years, gaps = extract_experience(text)
            tags = tag_skills(text)
            sem, hard = semantic_and_hard(jd_text, text)
            score = scoring_breakdown(jd_text, text, sem, hard, years, len(tags.get("soft", [])), social, contacts, edu, len(exp), custom_weights=st.session_state.custom_weights)
            eosint = email_osint(contacts.get("email"))
            facts = fact_check_name(contacts.get("name"))
            
            # Advanced features
            location = extract_location(text) if jd_location else None
            location_dist = calculate_location_distance(jd_location, location) if jd_location and location else None
            
            career_summary = generate_career_summary(exp, years)
            salary_est = estimate_salary(years, location, tags) if enable_salary_est else None
            
            github_data = None
            if enable_github and social.get("github"):
                github_data = check_github_activity(social["github"][0])
            
            hibp_data = None
            if enable_hibp and contacts.get("email"):
                hibp_data = check_hibp(contacts["email"])
            
            proficiency_map = {}
            if enable_proficiency:
                for cat, skills_list in tags.items():
                    for skill in skills_list:
                        proficiency_map[skill] = detect_skill_proficiency(text, skill)
            
            job_hopping = detect_job_hopping(exp) if enable_hopping else None
            
            # Modern AI-powered features
            entities = extract_entities_with_spacy(text) if enable_ner else {}
            readability = analyze_text_readability(text) if enable_readability else {}
            language = detect_language(text) if enable_lang_detect else "en"
            certifications = extract_certifications(text) if enable_cert_extract else []
            cv_quality_score = calculate_cv_quality_score(text, readability, entities) if enable_quality else 0.7
            
            results.append({
                "filename": f.name, "text": text, "metadata": meta, "contacts": contacts, "social": social,
                "education": edu, "experience": exp, "years_experience": years, "gaps": gaps, "tags": tags,
                "semantic": sem, "hard": hard, "score": score, "email_osint": eosint, "fact_links": facts,
                "location": location, "location_distance": location_dist, "career_summary": career_summary,
                "salary_estimate": salary_est, "github_activity": github_data, "hibp": hibp_data,
                "skill_proficiency": proficiency_map, "job_hopping": job_hopping,
                "entities": entities, "readability": readability, "language": language,
                "certifications": certifications, "cv_quality_score": cv_quality_score
            })
        
        # Clear progress indicators
        progress_bar.empty()
        status_text.empty()
        
        st.success(f"✅ Processed {total_files} candidates successfully!")
        
        results = deduplicate(results)
        
        # Apply blind screening if enabled
        if enable_blind:
            results = [blind_screening_filter(r) for r in results]
        
        # Store pre-filter count for tracking
        original_count = len(results)
        
        # Apply filters with live tracking
        if ft_query:
            results = [r for r in results if full_text_search(r["text"], ft_query)]
            if len(results) < original_count:
                st.info(f"🔍 Full-text filter: {len(results)}/{original_count} candidates match '{ft_query}'")
        if bool_query:
            before_bool = len(results)
            results = [r for r in results if boolean_search(r["text"], bool_query)]
            if len(results) < before_bool:
                st.info(f"🔍 Boolean filter: {len(results)}/{before_bool} candidates match '{bool_query}'")
        if min_years > 0:
            before_years = len(results)
            results = [r for r in results if r["years_experience"] >= min_years]
            if len(results) < before_years:
                st.info(f"📅 Experience filter: {len(results)}/{before_years} candidates have {min_years}+ years")
        
        # Location filter
        if jd_location:
            before_location = len(results)
            results = [r for r in results if r.get("location_distance") and r["location_distance"] <= 50]
            if len(results) < before_location:
                st.info(f"📍 Location filter: {len(results)}/{before_location} candidates within 50km of {jd_location}")
        
        # Apply score threshold
        if st.session_state.score_threshold > 0:
            before_threshold = len(results)
            results = [r for r in results if r["score"]["overall"] >= st.session_state.score_threshold]
            if len(results) < before_threshold:
                st.info(f"🎯 Score threshold: {len(results)}/{before_threshold} candidates score ≥ {st.session_state.score_threshold:.0%}")
        
        # Store in session for sidebar live stats
        st.session_state.last_results = results
        
        # Apply Quick Search from sidebar
        if st.session_state.quick_search:
            search_query = st.session_state.quick_search.lower()
            before_search = len(results)
            results = [r for r in results if 
                      search_query in r.get("contacts", {}).get("name", "").lower() or
                      search_query in r.get("contacts", {}).get("email", "").lower() or
                      search_query in r["text"].lower()]
            if len(results) < before_search:
                st.info(f"🔍 Quick search: {len(results)}/{before_search} candidates match '{st.session_state.quick_search}'")
        
        # Apply View Mode Filter from sidebar
        if st.session_state.view_mode != "All":
            before_view = len(results)
            if st.session_state.view_mode == "⭐ Bookmarked":
                results = [r for r in results if r["filename"] in st.session_state.bookmarked]
            elif st.session_state.view_mode == "🟢 Strong (≥80%)":
                results = [r for r in results if r["score"]["overall"] >= 0.80]
            elif st.session_state.view_mode == "🟡 Average (60-80%)":
                results = [r for r in results if 0.60 <= r["score"]["overall"] < 0.80]
            elif st.session_state.view_mode == "🔴 Weak (<60%)":
                results = [r for r in results if r["score"]["overall"] < 0.60]
            elif st.session_state.view_mode == "🌟 Shortlisted":
                results = [r for r in results if st.session_state.candidate_status.get(r["filename"]) == "shortlisted"]
            elif st.session_state.view_mode == "📞 Interview":
                results = [r for r in results if st.session_state.candidate_status.get(r["filename"]) == "interview"]
            elif st.session_state.view_mode == "❌ Rejected":
                results = [r for r in results if st.session_state.candidate_status.get(r["filename"]) == "rejected"]
            
            if len(results) < before_view:
                st.info(f"👁️ View mode: {len(results)}/{before_view} candidates in '{st.session_state.view_mode}'")
        
        if not results:
            st.warning(f"⚠️ No candidates match all active filters. Started with {original_count} candidates.")
        else:
            # Enhanced Live HUD Banner with Real-time Metrics
            st.markdown("---")
            hud_metric_cols = st.columns([1, 1, 1, 1, 1, 1])
            
            with hud_metric_cols[0]:
                total_count = len(results)
                st.metric("🎯 Showing", total_count, delta=f"of {original_count}")
            
            with hud_metric_cols[1]:
                strong_count = len([r for r in results if r["score"]["overall"] >= 0.80])
                st.metric("🟢 Strong", strong_count, delta=f"{(strong_count/total_count*100):.0f}%" if total_count > 0 else "0%")
            
            with hud_metric_cols[2]:
                avg_score = sum([r["score"]["overall"] for r in results]) / total_count if total_count > 0 else 0
                st.metric("📊 Avg Score", f"{avg_score:.0%}", delta="↑" if avg_score >= 0.70 else "↓", delta_color="normal" if avg_score >= 0.70 else "inverse")
            
            with hud_metric_cols[3]:
                bookmarked_count = len([r for r in results if r["filename"] in st.session_state.bookmarked])
                st.metric("⭐ Starred", bookmarked_count, delta=f"{(bookmarked_count/total_count*100):.0f}%" if total_count > 0 else "0%")
            
            with hud_metric_cols[4]:
                avg_years = sum([r["years_experience"] for r in results]) / total_count if total_count > 0 else 0
                st.metric("📅 Avg Exp", f"{avg_years:.1f}y", delta="↑" if avg_years >= 5 else "↓")
            
            with hud_metric_cols[5]:
                linkedin_count = len([r for r in results if len(r.get("social", {}).get("linkedin", [])) > 0])
                st.metric("💼 LinkedIn", linkedin_count, delta=f"{(linkedin_count/total_count*100):.0f}%" if total_count > 0 else "0%")
            
            # Live Score Distribution Mini Chart
            st.markdown("#### 📈 Live Score Distribution")
            score_ranges = {"🔴 <40%": 0, "🟠 40-60%": 0, "🟡 60-80%": 0, "🟢 80%+": 0}
            for r in results:
                score = r["score"]["overall"]
                if score < 0.40:
                    score_ranges["🔴 <40%"] += 1
                elif score < 0.60:
                    score_ranges["🟠 40-60%"] += 1
                elif score < 0.80:
                    score_ranges["🟡 60-80%"] += 1
                else:
                    score_ranges["🟢 80%+"] += 1
            
            dist_cols = st.columns(4)
            for idx, (label, count) in enumerate(score_ranges.items()):
                with dist_cols[idx]:
                    pct = (count / total_count * 100) if total_count > 0 else 0
                    st.metric(label, count, delta=f"{pct:.0f}%")
            
            st.markdown("---")
            st.success(generate_quick_stats(results))
            
            # Create tabs for different views
            tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
                "📊 Dashboard", "📋 Candidates", "🔬 Analytics", "🤖 AI Insights", "📈 Comparisons", "🎯 Skill Gaps", "👤 Details", "⭐ Shortlist"
            ])
            
            # Dashboard metrics
            metrics = create_dashboard_metrics(results)
            
            with tab1:
                st.subheader("📊 Executive Dashboard")
                
                # Active Filters Display
                active_filters = []
                if ft_query:
                    active_filters.append(f"🔍 Text: '{ft_query}'")
                if bool_query:
                    active_filters.append(f"🔍 Boolean: '{bool_query}'")
                if min_years > 0:
                    active_filters.append(f"📅 Years ≥ {min_years}")
                if jd_location:
                    active_filters.append(f"📍 Location: {jd_location}")
                if st.session_state.score_threshold > 0:
                    active_filters.append(f"🎯 Score ≥ {st.session_state.score_threshold:.0%}")
                if st.session_state.quick_search:
                    active_filters.append(f"🔍 Quick: '{st.session_state.quick_search}'")
                if st.session_state.view_mode != "All":
                    active_filters.append(f"👁️ View: {st.session_state.view_mode}")
                
                if active_filters:
                    st.info("**Active Filters:** " + " | ".join(active_filters))
                
                # Key metrics row
                col1, col2, col3, col4, col5 = st.columns(5)
                with col1:
                    delta_total = f"+{metrics['total'] - original_count}" if metrics['total'] < original_count else None
                    st.metric("Total Candidates", metrics["total"], delta=delta_total)
                with col2:
                    st.metric("Avg Score", f"{metrics['avg_score']}%", delta="High" if metrics['avg_score'] >= 70 else "Low")
                with col3:
                    st.metric("Avg Experience", f"{metrics['avg_years']} yrs", delta="↑" if metrics['avg_years'] >= 5 else "↓")
                with col4:
                    st.metric("LinkedIn %", f"{metrics['linkedin_pct']}%", delta="Good" if metrics['linkedin_pct'] >= 50 else "Low")
                with col5:
                    st.metric("GitHub %", f"{metrics['github_pct']}%", delta="Good" if metrics['github_pct'] >= 30 else "Low")
                
                # Score distribution
                col_a, col_b = st.columns(2)
                with col_a:
                    st.plotly_chart(create_score_distribution_chart(results), use_container_width=True)
                with col_b:
                    st.plotly_chart(create_experience_vs_score(results), use_container_width=True)
                
                # Distribution breakdown
                st.markdown("### 🎯 Score Distribution Breakdown")
                dist_col1, dist_col2, dist_col3, dist_col4 = st.columns(4)
                with dist_col1:
                    st.metric("Excellent (80%+)", metrics["score_distribution"]["excellent"], 
                             delta_color="normal")
                with dist_col2:
                    st.metric("Good (60-80%)", metrics["score_distribution"]["good"])
                with dist_col3:
                    st.metric("Average (40-60%)", metrics["score_distribution"]["average"])
                with dist_col4:
                    st.metric("Below Avg (<40%)", metrics["score_distribution"]["below"], 
                             delta_color="inverse")
                
                # Top skills
                st.markdown("### 🔥 Top 10 Skills in Pool")
                skills_cols = st.columns(5)
                for idx, (skill, count) in enumerate(metrics["top_skills"][:10]):
                    with skills_cols[idx % 5]:
                        st.metric(skill.title(), count)
            
            with tab2:
                st.subheader("📋 Ranked Candidates Table")
                
                # Enhanced table with quick actions
                for idx, r in enumerate(sorted(results, key=lambda x: (x["score"]["overall"], x["score"]["jd_fit"]), reverse=True)):
                    overall_pct = r["score"]["overall"] * 100
                    jd_fit_pct = r["score"]["jd_fit"] * 100
                    status_badge = get_status_badge(r["score"]["overall"], r["filename"])
                    bookmark_icon = "⭐" if r["filename"] in st.session_state.bookmarked else "☆"
                    
                    with st.expander(f"#{idx+1} {status_badge} {bookmark_icon} {r['filename']} - Overall: {overall_pct:.1f}% | JD Fit: {jd_fit_pct:.1f}%"):
                        col_info, col_actions = st.columns([3, 1])
                        
                        with col_info:
                            st.write(f"**Name:** {r['contacts'].get('name', 'N/A')}")
                            st.write(f"**Email:** {r['contacts'].get('email', 'N/A')}")
                            st.write(f"**Years:** {r['years_experience']:.1f} | **LinkedIn:** {len(r['social'].get('linkedin', []))} | **GitHub:** {len(r['social'].get('github', []))}")
                            
                            if r.get("salary_estimate"):
                                sal = r['salary_estimate']
                                st.write(f"**Salary Est:** ${sal['min']:,} - ${sal['max']:,} (Est: ${sal['estimated']:,})")
                            
                            if r.get("career_summary"):
                                st.info(f"💼 {r['career_summary']}")
                        
                        with col_actions:
                            st.markdown("**Quick Actions**")
                            
                            if st.button(bookmark_icon, key=f"bookmark_{idx}"):
                                toggle_bookmark(r["filename"])
                                st.rerun()
                            
                            status = st.selectbox("Status", ["new", "shortlisted", "interview", "rejected", "offer", "hired"], 
                                                 key=f"status_{idx}", 
                                                 index=["new", "shortlisted", "interview", "rejected", "offer", "hired"].index(
                                                     st.session_state.candidate_status.get(r["filename"], "new")))
                            if st.button("Update", key=f"update_{idx}"):
                                update_candidate_status(r["filename"], status)
                                st.success(f"Status updated to {status}")
                        
                        # Notes section
                        note_input = st.text_input("Add Note", key=f"note_{idx}")
                        if st.button("Save Note", key=f"save_note_{idx}") and note_input:
                            add_candidate_note(r["filename"], note_input)
                            st.success("Note added!")
                        
                        if r["filename"] in st.session_state.candidate_notes:
                            st.markdown("**📝 Notes:**")
                            for note in st.session_state.candidate_notes[r["filename"]]:
                                st.caption(f"{note['timestamp']}: {note['note']}")
                
                st.markdown("---")
                
                # Downloadable summary table
                df = pd.DataFrame([{
                    "Filename": r["filename"], "Name": r["contacts"].get("name"), "Email": r["contacts"].get("email"),
                    "JD Fit %": round(r["score"]["jd_fit"]*100, 2), "Overall %": round(r["score"]["overall"]*100, 2),
                    "Years": round(r["years_experience"], 1), "LinkedIn": len(r["social"].get("linkedin", [])),
                    "GitHub": len(r["social"].get("github", [])), 
                    "Status": st.session_state.candidate_status.get(r["filename"], "new"),
                    "Bookmarked": "Yes" if r["filename"] in st.session_state.bookmarked else "No"
                } for r in results]).sort_values(by=["Overall %", "JD Fit %"], ascending=False)
                
                st.dataframe(df, use_container_width=True)
                
                d1, d2 = st.columns(2)
                with d1:
                    st.download_button("Download JSON", data=json.dumps(results, default=str, indent=2), 
                                     file_name="vitasort_results.json")
                with d2:
                    st.download_button("Download CSV", data=df.to_csv(index=False), 
                                     file_name="vitasort_results.csv")
            
            with tab3:
                st.subheader("🔬 Advanced Analytics")
                
                col_x, col_y = st.columns(2)
                with col_x:
                    st.plotly_chart(create_skill_category_radar(results), use_container_width=True)
                with col_y:
                    st.plotly_chart(create_top_skills_bar(results), use_container_width=True)
                
                st.plotly_chart(create_timeline_gantt(results), use_container_width=True)
            
            with tab4:
                st.subheader("🤖 AI-Powered Insights")
                
                # Check if AI features are available
                has_ai_data = any(r.get("entities") or r.get("readability") or r.get("language") or r.get("certifications") for r in results)
                
                if not has_ai_data:
                    st.info("💡 Enable AI features in the Control Center (Features tab) to see advanced insights!")
                else:
                    # Top 10 candidates with AI insights
                    top_candidates = results[:10]
                    
                    for idx, r in enumerate(top_candidates, 1):
                        with st.expander(f"#{idx} {r['contacts'].get('name', 'Unknown')} — {r['filename']}", expanded=(idx <= 3)):
                            # Quality Score at top - FIXED LAYOUT
                            if r.get("cv_quality_score"):
                                quality_pct = r["cv_quality_score"] * 100
                                
                                # Use full width for gauge to prevent overflow
                                col_q1, col_q2 = st.columns([1, 2])
                                with col_q1:
                                    st.metric("📊 CV Quality", f"{quality_pct:.1f}%", 
                                            delta="High" if quality_pct >= 75 else ("Medium" if quality_pct >= 50 else "Low"))
                                    st.metric("🎯 Overall Score", f"{r['score']['overall']*100:.1f}%")
                                
                                with col_q2:
                                    # Quality gauge with fixed dimensions
                                    fig_gauge = go.Figure(go.Indicator(
                                        mode="gauge+number",
                                        value=quality_pct,
                                        domain={'x': [0, 1], 'y': [0, 1]},
                                        title={'text': "Quality Score", 'font': {'size': 14}},
                                        number={'font': {'size': 24}},
                                        gauge={
                                            'axis': {'range': [None, 100], 'tickfont': {'size': 10}},
                                            'bar': {'color': "darkblue", 'thickness': 0.7},
                                            'steps': [
                                                {'range': [0, 50], 'color': "#FFE5E5"},
                                                {'range': [50, 75], 'color': "#FFF4E5"},
                                                {'range': [75, 100], 'color': "#E5F5E5"}
                                            ],
                                            'threshold': {
                                                'line': {'color': "red", 'width': 3},
                                                'thickness': 0.75,
                                                'value': 70
                                            }
                                        }
                                    ))
                                    fig_gauge.update_layout(
                                        height=180, 
                                        margin=dict(l=10, r=10, t=40, b=10),
                                        paper_bgcolor="rgba(0,0,0,0)",
                                        font=dict(size=11)
                                    )
                                    st.plotly_chart(fig_gauge, use_container_width=True, key=f"gauge_{idx}_{r['filename']}")
                            
                            # Language & Certifications Row
                            if r.get("language") or r.get("certifications"):
                                col_l1, col_l2 = st.columns(2)
                                with col_l1:
                                    if r.get("language"):
                                        lang_name = {"en": "🇬🇧 English", "es": "🇪🇸 Spanish", "fr": "🇫🇷 French", 
                                                   "de": "🇩🇪 German", "pt": "🇵🇹 Portuguese", "zh-cn": "🇨🇳 Chinese"}.get(r["language"], f"🌐 {r['language'].upper()}")
                                        st.success(f"**Language:** {lang_name}")
                                with col_l2:
                                    if r.get("certifications"):
                                        st.success(f"**🎓 Certifications:** {len(r['certifications'])}")
                                        with st.expander("View Certifications"):
                                            for cert in r["certifications"]:
                                                st.write(f"• {cert}")
                            
                            # Named Entities
                            if r.get("entities") and any(r["entities"].values()):
                                st.markdown("#### 🏷️ Named Entities Extracted")
                                ent = r["entities"]
                                
                                col_e1, col_e2 = st.columns(2)
                                with col_e1:
                                    if ent.get("persons"):
                                        st.write("**👤 People:**")
                                        st.write(", ".join(ent["persons"][:10]))
                                    if ent.get("organizations"):
                                        st.write("**🏢 Organizations:**")
                                        st.write(", ".join(ent["organizations"][:10]))
                                with col_e2:
                                    if ent.get("locations"):
                                        st.write("**📍 Locations:**")
                                        st.write(", ".join(ent["locations"][:10]))
                                    if ent.get("dates"):
                                        st.write("**📅 Dates:**")
                                        st.write(", ".join(ent["dates"][:10]))
                                
                                if ent.get("technical_skills"):
                                    st.write("**💻 Technical Skills (NER):**")
                                    st.write(", ".join(ent["technical_skills"][:15]))
                            
                            # Readability Analysis
                            if r.get("readability"):
                                st.markdown("#### 📖 Readability Analysis")
                                read = r["readability"]
                                
                                col_r1, col_r2, col_r3, col_r4 = st.columns(4)
                                with col_r1:
                                    flesch = read.get("flesch_reading_ease", 0)
                                    flesch_label = "Very Easy" if flesch >= 80 else ("Easy" if flesch >= 60 else ("Medium" if flesch >= 40 else "Hard"))
                                    st.metric("Flesch Score", f"{flesch:.1f}", delta=flesch_label)
                                with col_r2:
                                    grade = read.get("flesch_kincaid_grade", 0)
                                    st.metric("Grade Level", f"{grade:.1f}")
                                with col_r3:
                                    st.metric("Words", f"{read.get('word_count', 0):,}")
                                with col_r4:
                                    st.metric("Reading Time", f"{read.get('reading_time', 0):.1f} min")
                                
                                # Additional readability metrics - MASSIVELY ENHANCED
                                with st.expander("📊 Advanced Readability Metrics (10+ Metrics)", expanded=False):
                                    st.markdown("**Grade Level Indices:**")
                                    col_adv1, col_adv2, col_adv3 = st.columns(3)
                                    with col_adv1:
                                        st.write(f"**SMOG Index:** {read.get('smog_index', 0):.1f}")
                                        st.write(f"**Coleman-Liau:** {read.get('coleman_liau', 0):.1f}")
                                        st.write(f"**Gunning Fog:** {read.get('gunning_fog', 0):.1f}")
                                    with col_adv2:
                                        st.write(f"**ARI Score:** {read.get('automated_readability_index', 0):.1f}")
                                        st.write(f"**Linsear Write:** {read.get('linsear_write', 0):.1f}")
                                        st.write(f"**Dale-Chall:** {read.get('dale_chall_score', 0):.1f}")
                                    with col_adv3:
                                        st.write(f"**Difficult Words:** {read.get('difficult_words', 0):,}")
                                        st.write(f"**Avg Words/Sentence:** {read.get('avg_words_per_sentence', 0):.1f}")
                                        st.write(f"**Avg Syllables/Word:** {read.get('avg_syllables_per_word', 0):.2f}")
                                    
                                    st.markdown("---")
                                    st.markdown("**Structural Analysis:**")
                                    col_struct1, col_struct2, col_struct3 = st.columns(3)
                                    with col_struct1:
                                        st.metric("Total Sentences", f"{read.get('sentence_count', 0):,}")
                                    with col_struct2:
                                        st.metric("Total Words", f"{read.get('word_count', 0):,}")
                                    with col_struct3:
                                        st.metric("Total Syllables", f"{read.get('syllable_count', 0):,}")
                                    
                                    # Readability interpretation
                                    st.markdown("---")
                                    st.markdown("**📈 Interpretation Guide:**")
                                    if flesch >= 80:
                                        st.success("✅ **Very Easy** - Suitable for all audiences (5th grade level)")
                                    elif flesch >= 60:
                                        st.info("✅ **Easy** - Conversational for most adults (8th-9th grade)")
                                    elif flesch >= 40:
                                        st.warning("⚠️ **Medium** - Requires college-level reading (10th-12th grade)")
                                    else:
                                        st.error("❌ **Hard** - Very difficult, requires advanced education (College+)")
                                    
                                    if read.get("error"):
                                        st.error(f"⚠️ Analysis Error: {read.get('error')}")
                    
                    # Enhanced Radar Chart for Top 5
                    st.markdown("---")
                    st.markdown("### 🎯 Multi-Dimensional Comparison (Top 5)")
                    if len(results) >= 2:
                        try:
                            radar_fig = create_enhanced_radar_chart(results)
                            st.plotly_chart(radar_fig, use_container_width=True)
                        except Exception as e:
                            st.warning(f"Could not generate enhanced radar chart: {e}")
            
            with tab5:
                st.subheader("📈 Top Candidates Comparison")
                comp_table = create_comparative_table(results)
                st.dataframe(comp_table, use_container_width=True)
                
                st.markdown("### 🎯 Head-to-Head Score Breakdown")
                if len(results) >= 2:
                    top_2 = results[:2]
                    comparison_cols = st.columns(2)
                    
                    for idx, r in enumerate(top_2):
                        with comparison_cols[idx]:
                            st.markdown(f"**{r['contacts'].get('name', 'Unknown')}**")
                            
                            breakdown = r["score"]["breakdown"]
                            fig = go.Figure(go.Bar(
                                x=list(breakdown.values()),
                                y=list(breakdown.keys()),
                                orientation='h',
                                marker_color=['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7', '#FFA07A']
                            ))
                            fig.update_layout(height=300, showlegend=False, 
                                            xaxis_title="Score", yaxis_title="Category")
                            st.plotly_chart(fig, use_container_width=True)
            
            with tab6:
                st.subheader("🎯 Skill Gap Analysis")
                st.markdown("**Red flags (❌) indicate missing JD skills**")
                gap_matrix = create_skill_gap_matrix(jd_text, results)
                if not gap_matrix.empty:
                    st.dataframe(gap_matrix, use_container_width=True)
                else:
                    st.info("Upload candidates to see skill gap analysis")
            
            with tab7:
                st.subheader("👤 Individual Candidate Details")
                for r in results:
                    with st.expander(f"{r['filename']} — {r['contacts'].get('name') or 'Unknown'}"):
                        c1, c2, c3 = st.columns(3)
                        with c1:
                            st.metric("JD Fit", f"{round(r['score']['jd_fit']*100, 2)}%")
                            st.metric("Overall", f"{round(r['score']['overall']*100, 2)}%")
                        with c2:
                            st.write("Breakdown")
                            st.json(r["score"]["breakdown"])
                        with c3:
                            st.write("Contacts")
                            st.json(r["contacts"])
                        with c3:
                            st.write("Contacts")
                            st.json(r["contacts"])
                        st.write("Social Links")
                        st.write({k: len(v) for k, v in r["social"].items()})
                        st.write("Education")
                        st.json(r["education"]) if r["education"] else st.write("—")
                        st.write("Experience Timeline (years≈{:.1f})".format(r["years_experience"]))
                        st.json(r["experience"]) if r["experience"] else st.write("—")
                        if r["gaps"]:
                            st.warning(f"Employment gaps detected: {len(r['gaps'])}")
                        if r["metadata"]:
                            st.write("PDF Metadata")
                            st.json(r["metadata"])  
                        if r["email_osint"]:
                            st.write("Email OSINT")
                            st.json(r["email_osint"])  
                        if r["fact_links"]:
                            st.write("Fact Check Links")
                            for url in r["fact_links"]:
                                st.write(url)
                        st.write("Skill Tags")
                        st.json(r["tags"]) if r["tags"] else st.write("—")
                        
                        # Advanced feature displays
                        if r.get("career_summary"):
                            st.write("📝 Career Summary")
                            st.info(r["career_summary"])
                        
                        if r.get("salary_estimate"):
                            st.write("💰 Salary Estimate")
                            sal = r["salary_estimate"]
                            st.success(f"${sal['min']:,} - ${sal['max']:,} (Est: ${sal['estimated']:,})")
                        
                        if r.get("location"):
                            st.write(f"📍 Location: {r['location']}")
                            if r.get("location_distance"):
                                st.write(f"Distance: {r['location_distance']:.1f} km")
                        
                        if r.get("github_activity"):
                            st.write("💻 GitHub Activity")
                            st.json(r["github_activity"])
                        
                        if r.get("hibp"):
                            st.write("🔐 Security Check (HIBP)")
                            if r["hibp"]["breached"]:
                                st.error(f"Email found in {r['hibp']['count']} breaches")
                            else:
                                st.success("No breaches found")
                        
                        if r.get("skill_proficiency"):
                            st.write("🎯 Skill Proficiency Levels")
                            prof_df = pd.DataFrame([{"Skill": k, "Level": v} for k, v in r["skill_proficiency"].items()])
                            st.dataframe(prof_df, use_container_width=True)
                        
                        if r.get("job_hopping"):
                            jh = r["job_hopping"]
                            st.write("🔄 Career Stability")
                            if jh.get("hopper"):
                                st.warning(f"Avg tenure: {jh['avg_tenure']} years across {jh['job_count']} jobs")
                            else:
                                st.success(f"Stable career: {jh['avg_tenure']} years avg tenure")
                        
                        jd_words = set(re.findall(r"\b\w+\b", jd_text.lower()))
                        cv_words = set(re.findall(r"\b\w+\b", r["text"].lower()))
                        missing = [w for w in jd_words if (w not in cv_words and len(w) > 3)][:50]
                        if missing:
                            st.info(f"JD keywords not found (sample): {', '.join(missing[:20])}")
            
            with tab8:
                st.subheader("⭐ Shortlist & Bookmarked Candidates")
                
                # Filter for shortlisted and bookmarked
                shortlisted = [r for r in results if st.session_state.candidate_status.get(r["filename"]) in ["shortlisted", "interview", "offer"]]
                bookmarked = [r for r in results if r["filename"] in st.session_state.bookmarked]
                
                col_s, col_b = st.columns(2)
                with col_s:
                    st.metric("Shortlisted/Interview/Offer", len(shortlisted))
                with col_b:
                    st.metric("Bookmarked", len(bookmarked))
                
                st.markdown("### 🌟 Shortlisted Candidates")
                if shortlisted:
                    short_df = pd.DataFrame([{
                        "Filename": r["filename"], "Name": r["contacts"].get("name"), "Email": r["contacts"].get("email"),
                        "Overall %": round(r["score"]["overall"]*100, 2), "Status": st.session_state.candidate_status.get(r["filename"], "new"),
                        "Years": round(r["years_experience"], 1)
                    } for r in shortlisted]).sort_values(by="Overall %", ascending=False)
                    st.dataframe(short_df, use_container_width=True)
                else:
                    st.info("No candidates shortlisted yet. Use the Candidates tab to mark candidates as shortlisted.")
                
                st.markdown("### ⭐ Bookmarked Candidates")
                if bookmarked:
                    book_df = pd.DataFrame([{
                        "Filename": r["filename"], "Name": r["contacts"].get("name"), "Email": r["contacts"].get("email"),
                        "Overall %": round(r["score"]["overall"]*100, 2), 
                        "Status": st.session_state.candidate_status.get(r["filename"], "new"),
                        "Years": round(r["years_experience"], 1)
                    } for r in bookmarked]).sort_values(by="Overall %", ascending=False)
                    st.dataframe(book_df, use_container_width=True)
                    
                    # Batch actions for bookmarked
                    st.markdown("### 🔧 Batch Actions")
                    batch_action = st.selectbox("Apply to all bookmarked", ["Select Action", "shortlisted", "interview", "rejected"])
                    if st.button("Apply Batch Action") and batch_action != "Select Action":
                        for r in bookmarked:
                            update_candidate_status(r["filename"], batch_action)
                        st.success(f"Updated {len(bookmarked)} candidates to '{batch_action}'")
                        st.rerun()
                else:
                    st.info("No candidates bookmarked yet. Click the ☆ icon in the Candidates tab to bookmark candidates.")


if mode == "Default":
    st.header("📝 Job Description (Default)")
    jd = st.text_area("Enter the job description", key="jd_classic")
    st.header("📄 Upload Resumes (PDF only)")
    up = st.file_uploader("Upload PDF files", type=["pdf"], accept_multiple_files=True, key="files_classic")
    if up and jd:
        try:
            resumes = []
            for file in up:
                if _PdfReader is None:
                    st.error("PDF processing library not found. Please install PyPDF2 or pypdf.")
                    st.stop()
                text = extract_text_from_pdf(file)
                resumes.append(text)
            scores = rank_resumes_tfidf(jd, resumes)
            results = pd.DataFrame({"🍎 Resume": [f.name for f in up], "🎯 VitaSort Score": scores, "📈 Ranking": range(1, len(scores)+1)}).sort_values(by="🎯 VitaSort Score", ascending=False)
            results["📈 Ranking"] = range(1, len(results) + 1)
            tab1, tab2, tab3, tab4, tab5 = st.tabs(["📊 Results Table", "🌟 Radar Analysis", "📊 Parallel Coordinates", "☁️ Word Clouds", "🔥 Skills Heatmap"])
            with tab1:
                st.subheader("📋 Detailed Rankings Table")
                st.dataframe(results, use_container_width=True)
                top_candidate = results.iloc[0]
                st.markdown(f"### 🏆 Top Match: {top_candidate['🍎 Resume']}")
                st.markdown(f"**VitaSort Score: {top_candidate['🎯 VitaSort Score']}/100**")
                c1, c2, c3, c4 = st.columns(4)
                with c1: st.metric("📊 Average Score", f"{scores.mean():.2f}")
                with c2: st.metric("🎯 Highest Score", f"{scores.max():.2f}")
                with c3: st.metric("📉 Lowest Score", f"{scores.min():.2f}")
                with c4: st.metric("📈 Score Range", f"{(scores.max() - scores.min()):.2f}")
            with tab2:
                st.subheader("🌟 Multi-Dimensional Radar Analysis")
                if len(results) >= 1:
                    radar_chart = create_radar_chart(results, resumes, jd)
                    st.plotly_chart(radar_chart, use_container_width=True)
                    st.markdown("• Technical • Data Science • Business • Communication • Experience")
                else:
                    st.info("Upload at least one resume to see radar analysis")
            with tab3:
                st.subheader("📊 Multi-Dimensional Parallel Coordinates")
                parallel_chart = create_parallel_coordinates(results, resumes, jd)
                st.plotly_chart(parallel_chart, use_container_width=True)
            with tab4:
                st.subheader("☁️ Word Cloud Analysis")
                if len(resumes) > 0:
                    top_resume_idx = results.index[0]
                    top_resume_text = resumes[top_resume_idx]
                    job_wc, resume_wc = create_word_cloud_comparison(jd, top_resume_text)
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown("**🎯 Job Description Keywords**")
                        fig, ax = plt.subplots(figsize=(8, 6))
                        ax.imshow(job_wc, interpolation='bilinear')
                        ax.axis('off')
                        st.pyplot(fig)
                    with col2:
                        st.markdown("**🏆 Top Resume Keywords**")
                        fig, ax = plt.subplots(figsize=(8, 6))
                        ax.imshow(resume_wc, interpolation='bilinear')
                        ax.axis('off')
                        st.pyplot(fig)
            with tab5:
                st.subheader("🔥 Skills Matching Analysis")
                skills_heatmap = create_skills_heatmap(jd, resumes, [f.name for f in up])
                st.plotly_chart(skills_heatmap, use_container_width=True)
        except Exception as e:
            st.error(f"🚨 Error: {str(e)}")
